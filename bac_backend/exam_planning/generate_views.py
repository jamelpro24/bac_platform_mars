import io
import re
import uuid
import random
import os
import copy
import zipfile
from datetime import datetime

from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from django.http import HttpResponse

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import fitz

from .models import Inscription, Serie, SurveillanceAssignment, ProfessorSchedule
from professeurs.models import Professeur

generated_docs = {}

BASE_DIR     = os.path.dirname(os.path.abspath(__file__))
MODELE_18    = os.path.join(BASE_DIR, 'modeles', 'modele1.docx')
MODELE_16    = os.path.join(BASE_DIR, 'modeles', 'modele2.docx')
MODELE_PRES  = os.path.join(BASE_DIR, 'modeles', 'model_pres_v5.docx')   # avec حضور
MODELE_PRES1 = os.path.join(BASE_DIR, 'modeles', 'model_pres1_v2.docx')  # sans حضور (porte)
MODELE_ENV   = os.path.join(BASE_DIR, 'modeles', 'modele_enveloppe_examen_v8.docx')
MODELE_REC    = os.path.join(BASE_DIR, 'modeles', 'modele_rec.docx')
MODELE_SORTIE = os.path.join(BASE_DIR, 'modeles', 'sortie_urgent_v2.docx')
MODELE_NUM   = os.path.join(BASE_DIR, '..', '..', 'numero', 'mod1.docx')
MODELE_AGENT = os.path.join(BASE_DIR, 'modeles', 'model_agent.docx')
MODELE_INVITATION = os.path.join(BASE_DIR, 'modeles', 'invitation_prof.docx')
MODELE_NUM   = os.path.abspath(MODELE_NUM)
FONT_PATH    = os.path.join(BASE_DIR, 'modeles', 'Tajawal-Bold.ttf')

JOURS_ARABE = {0:'الإثنين', 1:'الثلاثاء', 2:'الأربعاء', 3:'الخميس', 4:'الجمعة', 5:'السبت', 6:'الأحد'}
MOIS_ARABE  = {1:'جانفي', 2:'فيفري', 3:'مارس', 4:'أفريل', 5:'ماي', 6:'جوان',
               7:'جويلية', 8:'أوت', 9:'سبتمبر', 10:'أكتوبر', 11:'نوفمبر', 12:'ديسمبر'}

def format_date_arabic(date_str):
    try:
        d = datetime.strptime(date_str, '%Y-%m-%d')
        return f"يوم {JOURS_ARABE[d.weekday()]} {d.day:02d} {MOIS_ARABE[d.month]} {d.year}"
    except Exception:
        return date_str


def _format_time_arabic(t_start, t_end):
    """08:00-12:00 -> من 8 الى 12 | 10:00-11:30 -> من 10 الى 11 و 30"""
    if not t_start or not t_end:
        return ''
    s_h = int(t_start.strftime('%H'))
    s_m = int(t_start.strftime('%M'))
    e_h = int(t_end.strftime('%H'))
    e_m = int(t_end.strftime('%M'))
    if s_m == 0 and e_m == 0:
        return f"من {s_h} الى {e_h}"
    elif s_m == 0:
        return f"من {s_h} الى {e_h} و {e_m}"
    elif e_m == 0:
        return f"من {s_h} و {s_m} الى {e_h}"
    else:
        return f"من {s_h} و {s_m} الى {e_h} و {e_m}"


def replace_placeholder(doc, placeholder, value):
    """Replace placeholder in document paragraphs and tables"""
    try:
        from docx.oxml.ns import qn
        from copy import deepcopy
        
        pf_full = '{{' + placeholder + '}}'
        
        def replace_in_para(para):
            """Replace placeholders, handling split {{ across runs"""
            runs = list(para._element.iter(qn('w:r')))
            i = 0
            while i < len(runs):
                run = runs[i]
                t_elems = list(run.iter(qn('w:t')))
                
                for t in t_elems:
                    if not t.text:
                        continue
                    # Full placeholder in single t element
                    if pf_full in t.text:
                        t.text = t.text.replace(pf_full, value)
                        continue
                    
                    # Handle split: t ends with '{' (possibly after space)
                    if t.text.strip().endswith('{') and i + 2 < len(runs):
                        next_run = runs[i + 1]
                        next_t_elems = list(next_run.iter(qn('w:t')))
                        if next_t_elems and next_t_elems[0].text == placeholder:
                            t.text = ''
                            next_t_elems[0].text = value
                            # Clear closing }}
                            for j in range(i + 2, min(i + 5, len(runs))):
                                rj = runs[j]
                                for tj in rj.iter(qn('w:t')):
                                    if tj.text:
                                        tj.text = ''
                
                i += 1
        
        # Paragraphs
        for para in doc.paragraphs:
            replace_in_para(para)
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_para(para)
                             
    except Exception as e:
        print(f"Replace error: {e}")

def create_image_stream(numero, serie):
    width, height = 240, 60
    img  = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((0, 0, width, height), radius=15, fill="white", outline="black", width=4)
    try:
        font = ImageFont.truetype(FONT_PATH, size=60)
    except IOError:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), str(numero), font=font)
    x = (width  - (bbox[2] - bbox[0])) // 2
    y = (height - (bbox[3] - bbox[1])) // 2
    draw.text((x, y), str(numero), fill="black", font=font)
    stream = io.BytesIO()
    img.save(stream, format='PNG')
    stream.seek(0)
    return stream


def create_numero_image(numero):
    """Crée une image PNG avec le numéro dans un rectangle arrondi"""
    largeur_cm = 10
    hauteur_cm = 4
    dpi = 300
    largeur_px = int(largeur_cm * dpi / 2.54)
    hauteur_px = int(hauteur_cm * dpi / 2.54)
    
    img = Image.new("RGB", (largeur_px, hauteur_px), "white")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle([20, 20, largeur_px-40, hauteur_px-40], radius=40, fill="white", outline="black", width=10)
    
    # Ajouter logo si existe
    logo_path = os.path.join(BASE_DIR, '..', 'numero', 'icon.png')
    if os.path.exists(logo_path):
        logo = Image.open(logo_path).convert('RGBA')
        logo = logo.resize((largeur_px-20, hauteur_px-20))
        alpha = logo.split()[3]
        alpha = alpha.point(lambda p: int(p * 0.05))
        logo.putalpha(alpha)
        img.paste(logo, (0, 0), logo)
    
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype(FONT_PATH, size=300)
    
    # Centrer le texte
    bbox = draw.textbbox((0, 0), str(numero), font=font)
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    x = (largeur_px - text_width) // 2
    y = (hauteur_px - text_height) // 2
    draw.text((x, y), str(numero), fill=(0, 0, 0), font=font)
    
    stream = io.BytesIO()
    img.save(stream, format='PNG')
    stream.seek(0)
    return stream


def generate_numero_document(nums, serie_nom, salle_numero):
    """Génère un document avec les images de numéros — 1er numéro en (0,1), puis ligne par ligne"""
    if not os.path.exists(MODELE_NUM):
        raise FileNotFoundError(f"Modèle numéro introuvable : {MODELE_NUM}")
    
    doc = Document(MODELE_NUM)
    
    if not doc.tables:
        raise ValueError("Modèle sans table")
    
    tab = doc.tables[0]

    # Calculate image dimensions to fit cell without changing table size
    IMG_ASPECT = 10.0 / 4.0
    CELL_PAD = Inches(0.08)

    row_h_emu = tab.rows[1].height
    if not row_h_emu:
        row_h_emu = Inches(1.3)
    cell_max_h = row_h_emu - CELL_PAD
    col_w_emu = tab.columns[0].width
    cell_max_w = col_w_emu - CELL_PAD

    img_w = min(cell_max_h * IMG_ASPECT, cell_max_w)
    img_w = max(img_w, Inches(1))

    # Need rows: first num in cell(0,1), then 2 per row starting row 1
    rows_needed = max(0, (len(nums) - 1 + 1) // 2)  # rows after row 0
    existing_data_rows = len(tab.rows) - 1  # rows after row 0
    from copy import deepcopy
    # Remove extra template rows if too many
    while len(tab.rows) > rows_needed + 1:
        tbl = tab._element
        tbl.remove(tbl[-1])
    # Add rows if needed
    for _ in range(max(0, existing_data_rows), rows_needed):
        new_row = deepcopy(tab.rows[-1]._element)
        tab._element.append(new_row)

    def _fill_cell(cell, num):
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run()
        run.add_picture(create_numero_image(num), width=img_w)

    nb = 0
    # Cell (0,0) = salle + série centré
    c00 = tab.cell(0, 0)
    c00.paragraphs[0].clear()
    c00.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c00.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    r00 = c00.paragraphs[0].add_run(f"قاعة {salle_numero}\nسلسلة {serie_nom}")
    r00.font.name = 'Tajawal'
    r00.font.size = Pt(14)
    r00.font.bold = True

    if nb < len(nums):
        _fill_cell(tab.cell(0, 1), nums[nb])
        nb += 1

    for i in range(1, rows_needed + 1):
        for j in range(min(2, len(tab.columns))):
            if nb >= len(nums):
                break
            _fill_cell(tab.cell(i, j), nums[nb])
            nb += 1

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

# ==================== Plan de salle ====================

def fill_header(doc, data):
    replace_placeholder(doc, 'delegation', data.get('delegation', ''))
    replace_placeholder(doc, 'municipalité', data.get('municipalité', data.get('delegation', '')))

    tab = doc.tables[0]
    if len(tab.rows) > 0 and len(tab.rows[0].cells) > 0:
        cell = tab.cell(0, 0)
        run1 = cell.paragraphs[0].add_run(data['salle'][4:8])
        run1.font.name = 'Tajawal'
        run1.font.size = Pt(12)
        run1.font.bold = True

    tab = doc.tables[1]
    if len(tab.rows[0].cells) > 2:
        cell = tab.cell(0, 2)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(data['centre'])
        run.font.name = 'Tajawal'; run.font.size = Pt(12); run.font.bold = True

    if len(tab.rows[1].cells) > 1:
        cell = tab.cell(1, 1)
        run = cell.paragraphs[0].add_run(data['matiere'])
        run.font.name = 'Tajawal'; run.font.size = Pt(12); run.font.bold = True

    if len(tab.rows[2].cells) > 1:
        cell = tab.cell(2, 1)
        run = cell.paragraphs[0].add_run(data['section'])
        run.font.name = 'Tajawal'; run.font.size = Pt(12); run.font.bold = True
    if len(tab.rows[2].cells) > 3:
        cell = tab.cell(2, 3)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(data['serie']))
        run.font.name = 'Tajawal'; run.font.size = Pt(12); run.font.bold = True

    if len(tab.rows[3].cells) > 1:
        cell = tab.cell(3, 1)
        run = cell.paragraphs[0].add_run(format_date_arabic(data['date']))
        run.font.name = 'Tajawal'; run.font.size = Pt(12); run.font.bold = True

    return doc

def insert_images(doc, candidats_data, layout="18"):
    tables_indices = [2, 3, 4]
    slots = []
    for table_idx in tables_indices:
        if table_idx >= len(doc.tables): continue
        table = doc.tables[table_idx]
        for row in range(len(table.rows)):
            for col in [0, 3, 6]:
                if col + 1 >= len(table.rows[row].cells): continue
                slots.append((table_idx, row, col))

    random.shuffle(slots)
    for idx, candidat in enumerate(candidats_data):
        if idx >= len(slots): break
        num = candidat.get('num_ins') if isinstance(candidat, dict) else candidat
        serie_nom = candidat.get('serie_nom', '') if isinstance(candidat, dict) else ''
        table_idx, row, col = slots[idx]
        table = doc.tables[table_idx]
        cell_img = table.cell(row, col)
        cell_img.paragraphs[0].clear()
        img_stream = create_image_stream(num, serie_nom)
        run = cell_img.paragraphs[0].add_run()
        run.add_picture(img_stream, width=Inches(1.25))
        cell_serie = table.cell(row, col + 1)
        cell_serie.paragraphs[0].clear()
        run2 = cell_serie.paragraphs[0].add_run(serie_nom)
        run2.font.name = 'Tajawal'; run2.font.size = Pt(14); run2.font.bold = True
    return doc

def generate_one_document(data, nums, candidats_data=None):
    modele = MODELE_16 if str(data.get('layout', '18')) == '15' else MODELE_18
    if not os.path.exists(modele):
        raise FileNotFoundError(f"Fichier modèle introuvable : {modele}")
    doc = Document(modele)
    doc = fill_header(doc, data)
    if candidats_data:
        shuffled_candidats = candidats_data[:]
        random.shuffle(shuffled_candidats)
        doc = insert_images(doc, shuffled_candidats, data.get('layout', '18'))
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream, shuffled_candidats
    else:
        shuffled = nums[:]
        random.shuffle(shuffled)
        shuffled_dicts = [{'num_ins': n, 'serie_nom': data.get('serie', '')} for n in shuffled]
        doc = insert_images(doc, shuffled_dicts, data.get('layout', '18'))
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream, shuffled_dicts

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_documents(request):
    payload    = request.data
    matiere    = payload.get('matiere', '')
    section    = payload.get('section', '')
    date       = payload.get('date', '')
    centre     = request.user.centre or ''
    delegation = request.user.delegation or ''
    layout     = str(payload.get('layout', '18'))
    salles     = payload.get('salles', [])

    if not salles:
        return Response({"error": "aucune salle fournie"}, status=400)

    documents = []
    errors    = []

    # pre-fetch inscription pool per serie, shared across rooms
    serie_pool = {}
    for salle_data in salles:
        for serie_info in salle_data.get('series', []):
            sn = serie_info.get('nom')
            if sn and sn not in serie_pool:
                try:
                    so = Serie.objects.get(nom=sn, user=request.user)
                    serie_pool[sn] = list(Inscription.objects.filter(serie=so).order_by('num_ins').values_list('num_ins', flat=True))
                except Serie.DoesNotExist:
                    serie_pool[sn] = []

    for salle_data in salles:
        salle_numero = salle_data.get('salle_numero')
        section_str = salle_data.get('section_str', section)
        series_str = salle_data.get('series_str', '')
        nums_str = salle_data.get('nums_str', '')
        candidats = salle_data.get('candidats', [])
        
        if candidats:
            nums = [c.get('num_ins') for c in candidats if c.get('num_ins')]
            all_noms = [(c.get('num_ins'), c.get('nom_prenom')) for c in candidats if c.get('num_ins')]
        elif nums_str:
            nums = nums_str.split('-')
            all_noms = [(n, '') for n in nums if n]
        else:
            max_places = int(salle_data.get('layout', layout))
            for serie_info in salle_data.get('series', []):
                serie_nom = serie_info.get('nom')
                if not serie_nom: continue
                pool = serie_pool.get(serie_nom, [])
                if not pool:
                    errors.append(f"لا توجد أرقام متبقية للسلسلة {serie_nom}"); continue
                used = pool[:max_places]
                all_noms = [(n, '') for n in used]
                del pool[:max_places]
                nums = used
                series_str = serie_nom
                section_str = section or (Serie.objects.get(nom=serie_nom, user=request.user).section.nom if Serie.objects.filter(nom=serie_nom, user=request.user).exists() else section)
                break

        if not nums:
            errors.append(f"لا توجد أرقام للقاعة {salle_numero}"); continue

        max_places = int(salle_data.get('layout', layout))
        if len(nums) > max_places:
            errors.append(f"القاعة {salle_numero} : {len(nums)} مترشح، الحد الأقصى {max_places}"); continue
        room_layout = salle_data.get('layout', layout)
        data = {
            'centre': centre, 'delegation': delegation, 'municipalité': delegation, 'date': date,
            'matiere': matiere, 'section': section_str, 'serie': series_str,
            'salle': f"قاعة {salle_numero}", 'layout': room_layout,
        }
        
        try:
            stream, _ = generate_one_document(data, nums, candidats)
        except Exception as e:
            errors.append(f"Erreur génération salle {salle_numero} : {str(e)}"); continue
        doc_id = str(uuid.uuid4())
        label  = f"تصميم — {section_str} — قاعة {salle_numero} — {series_str} — {date}"
        generated_docs[doc_id] = {'stream': stream, 'serie': series_str, 'salle': salle_numero, 'label': label, 'type': 'plan', 'section': section_str}
        documents.append({'doc_id': doc_id, 'label': label, 'section': section_str, 'type': 'plan'})
            
        # ── 2. ملصق الظرف (envelope stickers) ──────────────────────────
        try:
            annee = request.user.annee_scolaire or ''
            delegation = request.user.delegation or ''
            centre = request.user.centre or ''
            env_data = {
                'delegation': delegation,
                'centre': centre,
                'matiere': matiere, 
                'section': section_str, 
                'serie': series_str,
                'salle': str(salle_numero), 
                'date': date, 
                'heure': '',
                'annee': annee,
                'nb_candidats': len(nums),
            }
            env_stream = _generate_envelope(env_data, nums)
            env_doc_id = str(uuid.uuid4())
            env_label = f"ملصق — {section_str} — قاعة {salle_numero} — {series_str}"
            generated_docs[env_doc_id] = {'stream': env_stream, 'serie': series_str, 'salle': salle_numero, 'label': env_label, 'type': 'envelope', 'section': section_str}
            documents.append({'doc_id': env_doc_id, 'label': env_label, 'section': section_str, 'type': 'envelope'})
        except Exception as e:
            errors.append(f"Erreur envelope salle {salle_numero} : {str(e)}")
        
        # ── 3. Document numéros (liste) ──────────────────────────
        try:
            num_stream = generate_numero_document(nums, series_str, str(salle_numero))
            num_doc_id = str(uuid.uuid4())
            num_label = f"أرقام — {section_str} — قاعة {salle_numero} — {series_str}"
            generated_docs[num_doc_id] = {'stream': num_stream, 'serie': series_str, 'salle': salle_numero, 'label': num_label, 'type': 'numero', 'section': section_str}
            documents.append({'doc_id': num_doc_id, 'label': num_label, 'section': section_str, 'type': 'numero'})
        except Exception as e:
            errors.append(f"Erreur numéros salle {salle_numero} : {str(e)}")

        # ── 4. Document sortie urgence ──────────────────────────
        try:
            inscriptions = [{'num_ins': n, 'nom_prenom': '', 'cin': ''} for n in nums]
            annee = request.user.annee_scolaire or ''
            sortie_data = {
                'delegation': delegation,
                'centre': centre,
                'annee_scolaire': annee,
                'section': section_str,
                'serie': series_str,
                'salle': str(salle_numero),
                'matiere': matiere,
                'date': format_date_arabic(date) if date else date,
            }
            sortie_stream = _generate_sortie(sortie_data, inscriptions)
            sortie_doc_id = str(uuid.uuid4())
            sortie_label = f"خروج — {section_str} — قاعة {salle_numero} — {series_str}"
            generated_docs[sortie_doc_id] = {'stream': sortie_stream, 'serie': series_str, 'salle': salle_numero, 'label': sortie_label, 'type': 'sortie', 'section': section_str}
            documents.append({'doc_id': sortie_doc_id, 'label': sortie_label, 'section': section_str, 'type': 'sortie'})
        except Exception as e:
            errors.append(f"Erreur sortie salle {salle_numero} : {str(e)}")
    if errors and not documents:
        return Response({"error": "\n".join(errors)}, status=400)

    documents = _merge_by_type(documents)

    response_data = {"documents": documents}
    if errors:
        response_data["warnings"] = errors
    return Response(response_data, status=200)


def _generate_one_invitation(schedules, prof, request, w, MODELE_INVITATION):
    """Generate a single invitation docx for one professor.
    schedules: list of ProfessorSchedule or SurveillanceAssignment objects
               (SurveillanceAssignment is auto-grouped by date)
    """
    try:
        import zipfile as zf
        from lxml import etree
        from copy import deepcopy

        centre = request.user.centre or ''
        delegation = request.user.delegation or ''
        annee_scolaire = request.user.annee_scolaire or ''
        nom_admin = request.user.nom_admin or ''
        annee = annee_scolaire.split('/')[-1] if '/' in annee_scolaire else annee_scolaire

        with zf.ZipFile(MODELE_INVITATION, 'r') as zin:
            template_items = {name: zin.read(name) for name in zin.namelist()}

        root = etree.fromstring(template_items['word/document.xml'])
        runs = list(root.iter(f'{{{w}}}r'))

        first = schedules[0]
        date_reu = format_date_arabic(str(first.date))

        repl = {
            'nom_prof': prof.nom,
            'nom_admin': nom_admin,
            'centre': centre,
            'annee_scolaire': annee_scolaire,
            'annee': annee,
            'date_reu': date_reu,
        }

        i = 0
        while i < len(runs):
            t_el = runs[i].find(f'{{{w}}}t')
            if t_el is None or not t_el.text:
                i += 1
                continue
            text = t_el.text

            handled = False
            for key, value in repl.items():
                pat = '{{' + key + '}}'
                if pat == text:
                    t_el.text = str(value)
                    handled = True
                    break
                elif pat in text:
                    t_el.text = text.replace(pat, str(value))
                    handled = True
                    break

            if handled:
                i += 1
                continue

            if text == '{{' and i + 2 < len(runs):
                mid_t = runs[i+1].find(f'{{{w}}}t')
                end_t = runs[i+2].find(f'{{{w}}}t')
                if mid_t is not None and end_t is not None and end_t.text == '}}' and mid_t.text in repl:
                    t_el.text = str(repl[mid_t.text])
                    mid_t.text = ''
                    end_t.text = ''
                    i += 3
                    continue

            # Handle {{key split across 2 runs: {{key in run N, }} in run N+1
            if text.startswith('{{') and not text.endswith('}}') and i + 1 < len(runs):
                key = text[2:]
                if key in repl:
                    next_t = runs[i+1].find(f'{{{w}}}t')
                    if next_t is not None and next_t.text and next_t.text.strip() == '}}':
                        t_el.text = str(repl[key])
                        next_t.text = ''
                        i += 2
                        continue

            i += 1

        all_tables = list(root.iter(f'{{{w}}}tbl'))

        for tbl in all_tables:
            rows = list(tbl.findall(f'{{{w}}}tr'))
            if len(rows) != 3:
                continue
            row_cell_counts = [len(list(r.iter(f'{{{w}}}tc'))) for r in rows]
            if row_cell_counts != [5, 5, 5]:
                continue

            # Build unified lookup: date -> {s1_type, s1_start, s1_end, s2_type, s2_start, s2_end}
            first = schedules[0]
            is_schedule = hasattr(first, 'session1_type')

            if is_schedule:
                dates_ordered = sorted(set(s.date for s in schedules))
                sched_lookup = {}
                for s in schedules:
                    sched_lookup[str(s.date)] = s
            else:
                # Convert SurveillanceAssignment list to date-keyed schedule
                from collections import defaultdict
                date_groups = defaultdict(lambda: {'s1_type': '', 's1_start': None, 's1_end': None,
                                                    's2_type': '', 's2_start': None, 's2_end': None})
                for a in schedules:
                    g = date_groups[str(a.date)]
                    if a.session_number == 1:
                        g['s1_type'] = a.type
                        g['s1_start'] = a.time_start
                        g['s1_end'] = a.time_end
                    elif a.session_number == 2:
                        g['s2_type'] = a.type
                        g['s2_start'] = a.time_start
                        g['s2_end'] = a.time_end
                dates_ordered = sorted(date_groups.keys())
                sched_lookup = date_groups

            session_days = dates_ordered[:4]

            for row_idx in range(3):
                row = rows[row_idx]
                cells = list(row.iter(f'{{{w}}}tc'))
                for ci in range(5):
                    if ci >= len(cells):
                        continue
                    cell = cells[ci]
                    if row_idx == 0:
                        if ci == 0:
                            _set_cell_text_xml(cell, "مادة", w)
                        elif ci - 1 < len(session_days):
                            _set_cell_text_xml(cell, format_date_arabic(str(session_days[ci - 1])), w)
                        else:
                            _set_cell_text_xml(cell, '', w)
                    else:
                        s = sched_lookup.get(str(session_days[ci - 1])) if ci - 1 < len(session_days) else None
                        if s is None:
                            _set_cell_text_xml(cell, '', w)
                            continue
                        if row_idx == 1:
                            if ci == 0:
                                _set_cell_text_xml(cell, "الحصة الأولى", w)
                            else:
                                s_start = s.session1_start if is_schedule else s['s1_start']
                                s_end = s.session1_end if is_schedule else s['s1_end']
                                _set_cell_text_xml(cell, _format_time_arabic(s_start, s_end) if s_start else '', w)
                        elif row_idx == 2:
                            if ci == 0:
                                _set_cell_text_xml(cell, "الحصة الثانية", w)
                            else:
                                s_start = s.session2_start if is_schedule else s['s2_start']
                                s_end = s.session2_end if is_schedule else s['s2_end']
                                _set_cell_text_xml(cell, _format_time_arabic(s_start, s_end) if s_start else '', w)

        docx_buf = io.BytesIO()
        with zf.ZipFile(MODELE_INVITATION, 'r') as zin:
            with zf.ZipFile(docx_buf, 'w', zf.ZIP_DEFLATED) as tmpz:
                for name in zin.namelist():
                    if name == 'word/document.xml':
                        tmpz.writestr(name, etree.tostring(root, xml_declaration=True, encoding='UTF-8'))
                    else:
                        tmpz.writestr(name, zin.read(name))

        docx_buf.seek(0)
        return docx_buf
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"_generate_one_invitation failed for prof {prof}: {e}")
        return None


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def download_invitations(request):
    """Generate invitation .docx for professors.
    - professeur_id: generate single docx for one professor
    - plan_id / plan_ids: filter by plan (fallback only)
    - Uses ProfessorSchedule if available, else falls back to SurveillanceAssignment
    """
    import zipfile as zf
    from collections import defaultdict

    w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    plan_id = request.data.get('plan_id')
    plan_ids = request.data.get('plan_ids')
    professeur_id = request.data.get('professeur_id')

    # Try ProfessorSchedule first (global confirmed snapshot, no plan_id needed)
    sched_qs = ProfessorSchedule.objects.filter(
        professeur__surv_assignments__plan__session__user=request.user
    ).distinct()
    if professeur_id:
        sched_qs = sched_qs.filter(professeur_id=professeur_id)

    sched_qs = sched_qs.select_related('professeur').order_by('professeur', 'date')
    sched_list = list(sched_qs)

    if sched_list:
        # Group by professor
        prof_schedules = defaultdict(list)
        prof_ids = set()
        for s in sched_list:
            prof_schedules[s.professeur_id].append(s)
            prof_ids.add(s.professeur_id)
        profs = {p.id: p for p in Professeur.objects.filter(id__in=prof_ids)}

        def _build_one(pid, schedules, prof):
            docx_buf = _generate_one_invitation(schedules, prof, request, w, MODELE_INVITATION)
            return docx_buf

        # Single professor
        if professeur_id:
            pid = int(professeur_id)
            schedules = prof_schedules.get(pid, [])
            if not schedules:
                return Response({"error": "لا توجد تعيينات لهذا الأستاذ"}, status=404)
            prof = profs.get(pid)
            if not prof:
                return Response({"error": "بيانات الأستاذ غير موجودة"}, status=404)
            docx_buf = _build_one(pid, schedules, prof)
            if docx_buf is None:
                return Response({"error": "فشل تحميل الدعوات"}, status=500)
            response = HttpResponse(
                docx_buf.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            safe_name = re.sub(r'[\\/*?:"<>|]', '_', prof.nom)
            response['Content-Disposition'] = f'attachment; filename="invitation_{safe_name}.docx"'
            return response

        # Multiple professors → ZIP
        zip_buf = io.BytesIO()
        errors = []
        doc_count = 0
        with zf.ZipFile(zip_buf, 'w', zf.ZIP_DEFLATED) as zout:
            for pid, schedules in prof_schedules.items():
                try:
                    prof = profs.get(pid)
                    if not prof:
                        errors.append(f"الأستاذ {pid} غير موجود")
                        continue
                    docx_buf = _build_one(pid, schedules, prof)
                    if docx_buf is None:
                        errors.append(f"فشل إنشاء دعوة للأستاذ {prof.nom}")
                        continue
                    safe_name = re.sub(r'[\\/*?:"<>|]', '_', prof.nom)
                    zout.writestr(f"invitation_{safe_name}.docx", docx_buf.read())
                    doc_count += 1
                except Exception as e:
                    errors.append(f"خطأ للأستاذ {pid}: {str(e)}")
                    continue

        if doc_count == 0:
            err_msg = "\n".join(errors) if errors else "فشل إنشاء جميع الدعوات"
            return Response({"error": err_msg}, status=500)

        zip_buf.seek(0)
        response = HttpResponse(zip_buf.getvalue(), content_type='application/zip')
        now_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        response['Content-Disposition'] = f'attachment; filename="invitations_profs_{now_str}.zip"'
        return response

    # Fallback: no ProfessorSchedule → use SurveillanceAssignment
    qs = SurveillanceAssignment.objects.filter(
        plan__session__user=request.user
    ).select_related('professeur')

    if plan_ids:
        qs = qs.filter(plan_id__in=plan_ids)
    elif plan_id:
        qs = qs.filter(plan_id=plan_id)
    if professeur_id:
        qs = qs.filter(professeur_id=professeur_id)

    qs = qs.order_by('professeur', 'date', 'session_number')

    prof_groups = defaultdict(list)
    for a in qs:
        prof_groups[a.professeur_id].append(a)

    profs = {p.id: p for p in Professeur.objects.filter(id__in=prof_groups.keys())}

    if not prof_groups:
        return Response({"error": "لا توجد تعيينات حراسة للأساتذة\nقم بتأكيد برنامج الأساتذة أولاً"}, status=404)

    # Single professor
    if professeur_id:
        pid = int(professeur_id)
        assigns = prof_groups.get(pid, [])
        if not assigns:
            return Response({"error": "لا توجد تعيينات لهذا الأستاذ"}, status=404)
        prof = profs.get(pid)
        if not prof:
            return Response({"error": "بيانات الأستاذ غير موجودة"}, status=404)
        docx_buf = _generate_one_invitation(assigns, prof, request, w, MODELE_INVITATION)
        if docx_buf is None:
            return Response({"error": "فشل تحميل الدعوات"}, status=500)
        response = HttpResponse(
            docx_buf.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        safe_name = re.sub(r'[\\/*?:"<>|]', '_', prof.nom)
        response['Content-Disposition'] = f'attachment; filename="invitation_{safe_name}.docx"'
        return response

    # Multiple professors → ZIP
    zip_buf = io.BytesIO()
    errors = []
    doc_count = 0
    with zf.ZipFile(zip_buf, 'w', zf.ZIP_DEFLATED) as zout:
        for pid, assigns in prof_groups.items():
            try:
                prof = profs.get(pid)
                if not prof:
                    errors.append(f"الأستاذ {pid} غير موجود")
                    continue
                docx_buf = _generate_one_invitation(assigns, prof, request, w, MODELE_INVITATION)
                if docx_buf is None:
                    errors.append(f"فشل إنشاء دعوة للأستاذ {prof.nom}")
                    continue
                safe_name = re.sub(r'[\\/*?:"<>|]', '_', prof.nom)
                zout.writestr(f"invitation_{safe_name}.docx", docx_buf.read())
                doc_count += 1
            except Exception as e:
                errors.append(f"خطأ للأستاذ {pid}: {str(e)}")
                continue

    if doc_count == 0:
        err_msg = "\n".join(errors) if errors else "فشل إنشاء جميع الدعوات"
        return Response({"error": err_msg}, status=500)

    zip_buf.seek(0)
    response = HttpResponse(zip_buf.getvalue(), content_type='application/zip')
    now_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="invitations_profs_{now_str}.zip"'
    return response


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def download_invitations_pdf(request):
    """Generate all invitations merged into a single PDF for printing."""
    import zipfile as zf
    from collections import defaultdict

    w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    plan_id = request.data.get('plan_id')
    plan_ids = request.data.get('plan_ids')

    # Try ProfessorSchedule first
    sched_qs = ProfessorSchedule.objects.filter(
        professeur__surv_assignments__plan__session__user=request.user
    ).distinct().select_related('professeur').order_by('professeur', 'date')
    sched_list = list(sched_qs)

    doc_streams = []

    if sched_list:
        prof_schedules = defaultdict(list)
        for s in sched_list:
            prof_schedules[s.professeur_id].append(s)
        profs = {p.id: p for p in Professeur.objects.filter(id__in=prof_schedules.keys())}

        for pid, schedules in prof_schedules.items():
            prof = profs.get(pid)
            if not prof:
                continue
            try:
                buf = _generate_one_invitation(schedules, prof, request, w, MODELE_INVITATION)
                if buf:
                    doc_streams.append(buf)
            except Exception:
                continue
    else:
        # Fallback: SurveillanceAssignment
        qs = SurveillanceAssignment.objects.filter(
            plan__session__user=request.user
        ).select_related('professeur')
        if plan_ids:
            qs = qs.filter(plan_id__in=plan_ids)
        elif plan_id:
            qs = qs.filter(plan_id=plan_id)
        qs = qs.order_by('professeur', 'date', 'session_number')

        prof_groups = defaultdict(list)
        for a in qs:
            prof_groups[a.professeur_id].append(a)
        profs = {p.id: p for p in Professeur.objects.filter(id__in=prof_groups.keys())}

        for pid, assigns in prof_groups.items():
            prof = profs.get(pid)
            if not prof:
                continue
            try:
                buf = _generate_one_invitation(assigns, prof, request, w, MODELE_INVITATION)
                if buf:
                    doc_streams.append(buf)
            except Exception:
                continue

    if not doc_streams:
        return Response({"error": "لا توجد دعوات لإنشاء PDF"}, status=404)

    pdf_stream = docs_to_pdf(doc_streams)

    response = HttpResponse(
        pdf_stream.read(),
        content_type='application/pdf',
    )
    now_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="invitations_profs_{now_str}.pdf"'
    return response


def _get_exam_info_for_assignment(assignment, user):
    """Get exam info (matiere) for a surveillance assignment."""
    try:
        from django.db.models import Q
        from .models import Examen, ExamenSalle, Salle
        from datetime import time as dt_time
        q = Q(examen__date=assignment.date) & Q(examen__user=user) & Q(salle__numero=assignment.salle_numero)
        if assignment.time_start:
            q &= Q(examen__heure_debut=assignment.time_start)
        es = ExamenSalle.objects.filter(q).select_related('examen__matiere').first()
        if es and es.examen and es.examen.matiere:
            return es.examen.matiere.nom
    except Exception:
        pass
    return assignment.type


def _set_cell_text_xml(cell, text, w):
    """Clear cell content and set simple text."""
    from lxml import etree
    for para in cell.iter(f'{{{w}}}p'):
        for run in para.iter(f'{{{w}}}r'):
            for t in run.iter(f'{{{w}}}t'):
                t.text = ''
    # Set text in first run of first paragraph
    paras = list(cell.iter(f'{{{w}}}p'))
    if paras:
        runs = list(paras[0].iter(f'{{{w}}}r'))
        if runs:
            t_el = runs[0].find(f'{{{w}}}t')
            if t_el is not None:
                t_el.text = str(text)
                return
            # Create t element
            t_el = etree.SubElement(runs[0], f'{{{w}}}t')
            t_el.text = str(text)
            return
        # Create run + t
        r_el = etree.SubElement(paras[0], f'{{{w}}}r')
        t_el = etree.SubElement(r_el, f'{{{w}}}t')
        t_el.text = str(text)


def docs_to_pdf(doc_streams):
    """Convertit une liste de streams docx en un seul PDF"""
    pdf_doc = fitz.open()
    for doc_stream in doc_streams:
        doc_stream.seek(0)
        docx_bytes = doc_stream.read()
        if not docx_bytes:
            continue
        try:
            docx_doc = fitz.open(stream=io.BytesIO(docx_bytes), filetype="docx")
            if docx_doc.page_count > 0:
                pdfbytes = docx_doc.convert_to_pdf()
                pdf_page = fitz.open(stream=io.BytesIO(pdfbytes), filetype="pdf")
                pdf_doc.insert_pdf(pdf_page)
                pdf_page.close()
            docx_doc.close()
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"docs_to_pdf: conversion failed: {e}")

    if pdf_doc.page_count == 0:
        pdf_doc.close()
        return io.BytesIO(b'%PDF-1.4\n%\xe2\xe3\xcf\xd3\n')

    output = io.BytesIO()
    pdf_doc.save(output)
    pdf_doc.close()
    output.seek(0)
    return output


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def download_combined_pdf(request):
    """Génère un PDF contenant les 5 documents (plan, envelope, numeros, presence, door) pour chaque salle"""
    doc_ids = request.data.get('doc_ids', [])
    filename = request.data.get('filename', 'documents')

    doc_streams = []
    for doc_id in doc_ids:
        doc_info = generated_docs.get(doc_id)
        if doc_info and doc_info.get('type') != 'sortie':
            doc_streams.append(doc_info['stream'])

    if not doc_streams:
        return Response({"error": "Aucun document à fusionner"}, status=400)

    pdf_stream = docs_to_pdf(doc_streams)

    response = HttpResponse(
        pdf_stream.read(),
        content_type='application/pdf',
    )
    response['Content-Disposition'] = f'inline; filename="{filename}.pdf"'
    return response


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_document(request, doc_id):
    doc_info = generated_docs.get(doc_id)
    if not doc_info:
        return Response({"error": "Document non trouvé"}, status=404)
    stream = doc_info['stream']
    stream.seek(0)
    response = HttpResponse(
        stream.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{doc_info.get("label", "document")}.docx"'
    return response


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def download_zip(request):
    doc_ids = request.data.get('doc_ids', [])
    filename = request.data.get('filename', 'documents')
    streams = []
    for doc_id in doc_ids:
        doc_info = generated_docs.get(doc_id)
        if doc_info:
            streams.append((doc_info['label'], doc_info['stream']))
    if not streams:
        return Response({"error": "Aucun document"}, status=400)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for label, stream in streams:
            stream.seek(0)
            zf.writestr(f"{label}.docx", stream.read())
    buf.seek(0)
    response = HttpResponse(buf.read(), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="{filename}.zip"'
    return response


def _generate_agent(data):
    """Génère document agent (modele_agent.docx)"""
    import zipfile
    from lxml import etree

    w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    with zipfile.ZipFile(MODELE_AGENT, 'r') as z:
        xml_bytes = z.read('word/document.xml')

    root = etree.fromstring(xml_bytes)

    replacements = {
        'delegation': data.get('delegation', ''),
        'annee_scolaire': data.get('annee', data.get('annee_scolaire', '')),
        'centre': data.get('centre', ''),
        'agent': data.get('agent', ''),
        'agent2': data.get('agent2', ''),
        'fonction': data.get('fonction', ''),
        'date_reu': data.get('date_reu', ''),
        'temp': data.get('temp', ''),
        'nom_admin': data.get('nom_admin', ''),
    }

    runs = list(root.iter(f'{{{w}}}r'))

    agent_names = [data.get('agent', ''), data.get('agent2', '')]
    agent_idx = 0

    i = 0
    while i < len(runs):
        t_el = runs[i].find(f'{{{w}}}t')
        if t_el is None or not t_el.text:
            i += 1
            continue
        text = t_el.text

        # Pattern 1: {{key}} in a single run
        handled = False
        for key, value in replacements.items():
            pat = '{{' + key + '}}'
            if pat == text:
                if key == 'agent':
                    t_el.text = str(agent_names[min(agent_idx, len(agent_names) - 1)])
                    agent_idx += 1
                elif key == 'agent2':
                    continue
                else:
                    t_el.text = str(value)
                handled = True
                break
            elif pat in text:
                t_el.text = text.replace(pat, str(value))
                handled = True
                break

        if handled:
            i += 1
            continue

        # Pattern 2: {{ + key + }}  (three runs — split)
        if text == '{{' and i + 2 < len(runs):
            mid_t = runs[i+1].find(f'{{{w}}}t')
            end_t = runs[i+2].find(f'{{{w}}}t')
            if mid_t is not None and end_t is not None and end_t.text == '}}' and mid_t.text in replacements:
                t_el.text = str(replacements[mid_t.text])
                mid_t.text = ''
                end_t.text = ''
                i += 3
                continue

        i += 1

    output = io.BytesIO()
    with zipfile.ZipFile(MODELE_AGENT, 'r') as z_in:
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as z_out:
            for item in z_in.namelist():
                if item == 'word/document.xml':
                    z_out.writestr(item, etree.tostring(root, xml_declaration=True, encoding='UTF-8'))
                else:
                    z_out.writestr(item, z_in.read(item))

    output.seek(0)
    return output


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_agent_document(request):
    payload = request.data
    data = {
        'delegation': request.user.delegation or '',
        'centre': payload.get('centre', request.user.centre or ''),
        'annee': request.user.annee_scolaire or '',
        'agent': payload.get('agent', ''),
        'agent2': payload.get('agent2', ''),
        'fonction': payload.get('fonction', ''),
        'date_reu': payload.get('date_reu', ''),
        'temp': payload.get('temp', ''),
        'nom_admin': payload.get('nom_admin', request.user.nom_admin or ''),
    }
    try:
        stream = _generate_agent(data)
    except Exception as e:
        return Response({"error": f"Erreur de génération : {str(e)}"}, status=400)

    doc_id = str(uuid.uuid4())
    label = f"أعوان الكتابة والتنسيق — {data['date_reu']}"
    generated_docs[doc_id] = {'stream': stream, 'label': label, 'type': 'agent'}
    return Response({"doc_id": doc_id, "label": label, "type": "agent"})


# ==================== HELPERS بطاقات الحضور / سجل القاعة ====================

def set_cell_text(cell, text, font_size_pt=12, bold=True, font_name='Tajawal', align=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from lxml import etree
    
    para = cell.paragraphs[0]
    para.clear()
    
    para.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    
    # No wrap - text in single line
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn('w:tcPr'))
    
    noWrap = tcPr.find(qn('w:noWrap'))
    if noWrap is None:
        noWrap = etree.SubElement(tcPr, qn('w:noWrap'))
    noWrap.set(qn('w:val'), '1')


def _delete_row(table, row_idx):
    """Supprime la ligne à l'index donné."""
    table._tbl.remove(table.rows[row_idx]._tr)


def _clone_row(table, source_row_idx):
    """Clone la ligne source et l'ajoute à la fin du tableau. Retourne la nouvelle ligne."""
    new_tr = copy.deepcopy(table.rows[source_row_idx]._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _clear_row_cells(row):
    """Vide le texte de toutes les cellules."""
    for cell in row.cells:
        for para in cell.paragraphs:
            para.clear()


def _replace_placeholder_simple(doc, placeholder, value):
    """Replace placeholder by searching all XML elements"""
    from docx.oxml.ns import qn
    
    value = str(value)
    pattern = '{{' + placeholder + '}}'
    
    # Find all w:t elements (text elements)
    for elem in doc.element.body.iter():
        if elem.tag == qn('w:t') and elem.text and pattern in elem.text:
            elem.text = elem.text.replace(pattern, value)
    
    # Also check in shapes/textboxes (wsp:txbx)
    for elem in doc.element.body.iter(qn('w:txbx')):
        for t in elem.iter(qn('w:t')):
            if t.text and pattern in t.text:
                t.text = t.text.replace(pattern, value)
    
    # Check alternate content
    for elem in doc.element.body.iter(qn('mc:AlternateContent')):
        for t in elem.iter(qn('w:t')):
            if t.text and pattern in t.text:
                t.text = t.text.replace(pattern, value)
                print(f"[ENVELOPE]   FOUND in w:t: '{elem.text}'")
    
    print(f"[ENVELOPE] Total '{pattern}' found: {count}")
    
    # Also check in shapes/textboxes (wsp:txbx)
    for elem in doc.element.body.iter(qn('w:txbx')):
        for t in elem.iter(qn('w:t')):
            if t.text and pattern in t.text:
                t.text = t.text.replace(pattern, value)
                count += 1
                print(f"[ENVELOPE]   FOUND in txbx: '{t.text}'")
    
    # Check alternate content
    for elem in doc.element.body.iter(qn('mc:AlternateContent')):
        for t in elem.iter(qn('w:t')):
            if t.text and pattern in t.text:
                t.text = t.text.replace(pattern, value)
                count += 1
                print(f"[ENVELOPE]   FOUND in AlternateContent: '{t.text}'")


def _generate_envelope(data, nums):
    """Génère ملصق الظرف (stickers pour enveloppes)"""
    import zipfile
    from lxml import etree
    
    w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    with zipfile.ZipFile(MODELE_ENV, 'r') as z:
        xml_bytes = z.read('word/document.xml')
    
    root = etree.fromstring(xml_bytes)
    
    replacements = {
        'delegation': data.get('delegation', ''),
        'annee_scolaire': data.get('annee', data.get('annee_scolaire', '')),
        'centre': data.get('centre', ''),
        'section': data.get('section', ''),
        'serie': str(data.get('serie', '')),
        's': str(data.get('salle', '')),
        'nb': str(data.get('nb_candidats', 0)),
        'date': data.get('date', ''),
        'matiere': data.get('matiere', ''),
    }
    
    # Collect all w:r elements (text runs)
    runs = list(root.iter(f'{{{w}}}r'))
    
    i = 0
    while i < len(runs):
        t_el = runs[i].find(f'{{{w}}}t')
        if t_el is None or not t_el.text:
            i += 1
            continue
        text = t_el.text
        
        # Pattern 1: {{key}} in a single run
        handled = False
        for key, value in replacements.items():
            pat = '{{' + key + '}}'
            if pat == text:
                t_el.text = str(value)
                handled = True
                break
            elif pat in text:
                t_el.text = text.replace(pat, str(value))
                handled = True
                break
        
        if handled:
            i += 1
            continue
        
        # Pattern 2: { + { + key + }}  (four runs — for serie)
        if text == '{' and i + 3 < len(runs):
            r2_t = runs[i+1].find(f'{{{w}}}t')
            r3_t = runs[i+2].find(f'{{{w}}}t')
            r4_t = runs[i+3].find(f'{{{w}}}t')
            if (r2_t is not None and r2_t.text == '{' and
                r3_t is not None and r4_t is not None and r4_t.text == '}}' and
                r3_t.text in replacements):
                t_el.text = str(replacements[r3_t.text])
                r2_t.text = ''
                r3_t.text = ''
                r4_t.text = ''
                i += 4
                continue
        
        # Pattern 3: { + {key}}  (two runs — for nb)
        if text == '{' and i + 1 < len(runs):
            next_t = runs[i+1].find(f'{{{w}}}t')
            if next_t is not None:
                # Check if next text matches {key}
                for key, value in replacements.items():
                    pat = '{' + key + '}}'
                    if next_t.text == pat:
                        t_el.text = str(value)
                        next_t.text = ''
                        i += 2
                        handled = True
                        break
            if handled:
                continue
        
        # Pattern 4: {{ + key + }}  (three runs — for delegation, annee_scolaire)
        if text == '{{' and i + 2 < len(runs):
            mid_t = runs[i+1].find(f'{{{w}}}t')
            end_t = runs[i+2].find(f'{{{w}}}t')
            if mid_t is not None and end_t is not None and end_t.text == '}}' and mid_t.text in replacements:
                t_el.text = str(replacements[mid_t.text])
                mid_t.text = ''
                end_t.text = ''
                i += 3
                continue
        
        i += 1
    
    output = io.BytesIO()
    with zipfile.ZipFile(MODELE_ENV, 'r') as z_in:
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as z_out:
            for item in z_in.namelist():
                if item == 'word/document.xml':
                    z_out.writestr(item, etree.tostring(root, xml_declaration=True, encoding='UTF-8'))
                else:
                    z_out.writestr(item, z_in.read(item))
    
    output.seek(0)
    return output


def _generate_sortie(data, inscriptions):
    """Génère sortie urgence"""
    doc = Document(MODELE_SORTIE)

    for key in ['delegation', 'centre', 'salle', 'date', 'session', 'section', 'matiere', 'serie']:
        replace_placeholder(doc, key, data.get(key, ''))

    # Add time if available
    if data.get('heure_debut'):
        replace_placeholder(doc, 'heure_debut', data.get('heure_debut', ''))
    if data.get('heure_fin'):
        replace_placeholder(doc, 'heure_fin', data.get('heure_fin', ''))

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def insert_envelope_nums(doc, numbers):
    """Insere les numeros dans le document"""
    # Chercher la premiere table avec des cellules
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 1:
            continue
        # Remplir les premieres cellules avec les numeros
        cell_idx = 0
        for num in numbers[:20]:  # Max 20 numeros
            if cell_idx >= len(table.rows) * len(table.rows[0].cells):
                break
            row = cell_idx // 4
            col = cell_idx % 4
            if row < len(table.rows) and col < len(table.rows[row].cells):
                cell = table.rows[row].cells[col]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(str(num))
                run.font.size = Pt(14)
                run.bold = True
            cell_idx += 1
        break
    return doc


def replace_placeholder_docx(doc, placeholder, value):
    """Replace placeholder simpler pour envelopes"""
    full_tag = '{{ ' + placeholder + ' }}'
    for para in doc.paragraphs:
        if full_tag in para.text:
            para.text = para.text.replace(full_tag, value)
    return doc


def fill_presence_header(doc, data):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Remplacer tous les placeholders
    replace_placeholder(doc, 'annee_scolaire', data.get('annee_scolaire', ''))
    replace_placeholder(doc, 'centre', data.get('centre', ''))
    replace_placeholder(doc, 'section', data.get('section', ''))
    replace_placeholder(doc, 'serie', data.get('serie', ''))
    replace_placeholder(doc, 'salle', data.get('salle', ''))
    replace_placeholder(doc, 'matiere', data.get('matiere', ''))
    replace_placeholder(doc, 'date', format_date_arabic(data.get('date', '')))
    replace_placeholder(doc, 'heure', data.get('heure', ''))


def build_presence_table(doc, inscriptions, include_hudur_col: bool, user):
    """
    Construit le tableau candidats dynamiquement.
    Nombre de lignes = nombre réel d'inscrits.
    Colonne معهد : lu depuis Inscription.etablissement (rempli à l'import depuis القائمة الاسمية).
    Table fits within page boundaries.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Cm, Twips
    from lxml import etree

    tab = doc.tables[1]
    n   = len(inscriptions)

    # Supprimer toutes les lignes data sauf la première (template)
    while len(tab.rows) > 2:
        _delete_row(tab, len(tab.rows) - 1)

    # Set table width to page width
    for row in tab.rows:
        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = etree.SubElement(tc, qn('w:tcPr'))
            
            # Set cell width to proportional value
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = etree.SubElement(tcPr, qn('w:tcW'))
            
            # Distribute width: col0=3%, col1=10%, col2=20%, col3=12%, col4=10%, col5=25%, col6=20%
            widths = [3, 10, 20, 12, 10, 25, 20] if include_hudur_col and len(row.cells) > 6 else [4, 12, 24, 14, 12, 34]
            w = widths[col_idx] if col_idx < len(widths) else 10
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'pct')

    for i in range(n):
        if i == 0:
            row = tab.rows[1]
            _clear_row_cells(row)
        else:
            row = _clone_row(tab, 1)
            # Apply same width settings to new row
            for col_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = etree.SubElement(tc, qn('w:tcPr'))
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = etree.SubElement(tcPr, qn('w:tcW'))
                widths = [3, 10, 20, 12, 10, 25, 20] if include_hudur_col and len(row.cells) > 6 else [4, 12, 24, 14, 12, 34]
                w = widths[col_idx] if col_idx < len(widths) else 10
                tcW.set(qn('w:w'), str(w))
                tcW.set(qn('w:type'), 'pct')

        insc = inscriptions[i]

        # institute from inscription (filled from القائمة الاسمية)
        etablissement = str(insc.get('etablissement', '') or '').strip()

        # col 0 : ع/ر
        set_cell_text(row.cells[0], str(i + 1),
                      font_size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        # col 1 : رقم المترشح
        set_cell_text(row.cells[1], str(insc.get('num_ins', '')),
                      font_size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        # col 2 : الاسم و اللقب
        set_cell_text(row.cells[2], str(insc.get('nom_prenom', '')),
                      font_size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        # col 3 : رقم بطاقة التعريف الوطنية
        set_cell_text(row.cells[3], str(insc.get('cin', '') or ''),
                      font_size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        # col 4 : الشعبة
        set_cell_text(row.cells[4], str(insc.get('section', '')),
                      font_size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        # col 5 : المعهد
        set_cell_text(row.cells[5], etablissement,
                      font_size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        # col 6 : الحضور — uniquement model_pres.docx
        if include_hudur_col and len(row.cells) > 6:
            set_cell_text(row.cells[6], '',
                          font_size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


def generate_presence_document(data, inscriptions, modele_path, include_hudur_col, user):
    if not os.path.exists(modele_path):
        raise FileNotFoundError(f"Modèle introuvable : {modele_path}")
    doc = Document(modele_path)
    fill_presence_header(doc, data)
    build_presence_table(doc, inscriptions, include_hudur_col, user)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


# ==================== DOCUMENT قائمة المترشحين (modele_rec) ====================

def build_rec_table(doc, inscriptions):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from copy import deepcopy

    tab = doc.tables[1]  # index 0 = info table, index 1 = candidate table
    n = len(inscriptions)

    # Remove all data rows except the first (index 1 = template row)
    while len(tab.rows) > 2:
        _delete_row(tab, len(tab.rows) - 1)

    for i in range(n):
        if i == 0:
            row = tab.rows[1]
            for cell in row.cells:
                cell.paragraphs[0].clear()
        else:
            row = _clone_row(tab, 1)

        cells = row.cells
        set_cell_text(cells[0], str(i + 1), font_size_pt=10)  # tally
        set_cell_text(cells[1], str(inscriptions[i].get('num_ins', '')), font_size_pt=10)
        set_cell_text(cells[2], str(inscriptions[i].get('nom_prenom', '')), font_size_pt=10)
        set_cell_text(cells[3], '', font_size_pt=10)  # signature
        set_cell_text(cells[4], '', font_size_pt=10)  # notes

def generate_rec_document(data, inscriptions):
    if not os.path.exists(MODELE_REC):
        raise FileNotFoundError(f"Modèle rec introuvable : {MODELE_REC}")
    doc = Document(MODELE_REC)
    for key in ['delegation', 'centre', 'salle', 'matiere', 'section', 'serie', 'date', 'annee_scolaire']:
        replace_placeholder(doc, key, data.get(key, ''))
    build_rec_table(doc, inscriptions)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


# ==================== VIEW بطاقات الحضور + سجل القاعة ====================

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_presence(request):
    """
    Génère pour chaque salle (avec candidates depuis frontend):
      • بطاقة الحضور  — model_pres.docx  (avec colonne حضور)
      • سجل القاعة    — model_pres1.docx (sans  colonne حضور — pour coller sur la porte)
    """
    payload = request.data
    matiere = payload.get('matiere', '')
    date    = payload.get('date',    '')
    heure   = payload.get('heure',   '')
    centre  = request.user.centre         or ''
    annee   = request.user.annee_scolaire or ''
    salles  = payload.get('salles',  [])

    if not salles:
        return Response({"error": "aucune salle fournie"}, status=400)

    documents = []
    errors    = []

    # pre-fetch candidate pool per serie, shared across rooms
    serie_cand_pool = {}
    for salle_data in salles:
        for serie_info in salle_data.get('series', []):
            sn = serie_info.get('nom')
            if sn and sn not in serie_cand_pool:
                try:
                    so = Serie.objects.get(nom=sn, user=request.user)
                    inscs = Inscription.objects.filter(serie=so)
                    serie_cand_pool[sn] = [{
                        'num_ins': ins.num_ins,
                        'nom_prenom': ins.nom_prenom,
                        'cin': ins.cin or '',
                        'section_nom': ins.section or '',
                        'etablissement': ins.etablissement or '',
                    } for ins in inscs]
                except Serie.DoesNotExist:
                    serie_cand_pool[sn] = []

    for salle_data in salles:
        salle_numero = salle_data.get('salle_numero')
        section_str = salle_data.get('section_str', '')
        series_str = salle_data.get('series_str', '')
        candidats = salle_data.get('candidats', [])
        
        if not candidats:
            for serie_info in salle_data.get('series', []):
                serie_nom = serie_info.get('nom')
                if not serie_nom: continue
                pool = serie_cand_pool.get(serie_nom, [])
                if not pool:
                    errors.append(f"لا توجد أرقام متبقية للسلسلة {serie_nom}"); continue
                room_cap = int(salle_data.get('layout', 18))
                candidats = pool[:room_cap]
                del pool[:room_cap]
                if not series_str:
                    series_str = serie_nom
                if not section_str:
                    try:
                        so = Serie.objects.get(nom=serie_nom, user=request.user)
                        section_str = so.section.nom if so.section else section_str
                    except Serie.DoesNotExist:
                        pass
                break
            if not candidats:
                errors.append(f"Aucun candidat pour la salle {salle_numero}"); continue

        inscriptions = [{
            'num_ins': c.get('num_ins', ''),
            'nom_prenom': c.get('nom_prenom', ''),
            'cin': c.get('cin') or c.get('cin', '') or '',
            'section': c.get('section_nom', ''),
            'etablissement': c.get('etablissement', ''),
        } for c in candidats]

        etablissement = candidats[0].get('etablissement', '') if candidats else ''
        data = {
            'centre':          centre,
            'annee_scolaire':  annee,
            'etablissement':   etablissement,
            'section':         section_str,
            'serie':           series_str,
            'salle':           str(salle_numero),
            'matiere':         matiere,
            'date':            date,
            'heure':           heure,
        }

        # ── 1. بطاقة الحضور (avec حضور) ──────────────────────────
        try:
            stream = generate_presence_document(
                data, inscriptions,
                modele_path=MODELE_PRES,
                include_hudur_col=True,
                user=request.user,
            )
            doc_id = str(uuid.uuid4())
            label  = f"حضور — {section_str} — قاعة {salle_numero} — {series_str} — {date}"
            generated_docs[doc_id] = {'stream': stream, 'label': label, 'section': section_str, 'type': 'presence'}
            documents.append({'doc_id': doc_id, 'label': label, 'section': section_str, 'type': 'presence'})
        except Exception as e:
            errors.append(f"بطاقة حضور salle {salle_numero} : {str(e)}")

        # ── 2. سجل القاعة (sans حضور — porte) ────────────────────
        try:
            stream2 = generate_presence_document(
                data, inscriptions,
                modele_path=MODELE_PRES1,
                include_hudur_col=False,
                user=request.user,
            )
            doc_id2 = str(uuid.uuid4())
            label2  = f"سجل القاعة — {section_str} — قاعة {salle_numero} — {series_str} — {date}"
            generated_docs[doc_id2] = {'stream': stream2, 'label': label2, 'section': section_str, 'type': 'door'}
            documents.append({'doc_id': doc_id2, 'label': label2, 'section': section_str, 'type': 'door'})
        except Exception as e:
            errors.append(f"سجل القاعة salle {salle_numero} : {str(e)}")

    if errors and not documents:
        return Response({"error": "\n".join(errors)}, status=400)

    documents = _merge_by_type(documents)

    response_data = {"documents": documents}
    if errors:
        response_data["warnings"] = errors
    return Response(response_data, status=200)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_rec(request):
    """
    Génère la liste des élèves (modele_rec.docx) pour chaque salle.
    """
    payload = request.data
    matiere = payload.get('matiere', '')
    date    = payload.get('date', '')
    centre  = request.user.centre or ''
    annee   = request.user.annee_scolaire or ''
    salles  = payload.get('salles', [])

    if not salles:
        return Response({"error": "aucune salle fournie"}, status=400)

    documents = []
    errors = []

    # pre-fetch candidate pool per serie
    rec_pool = {}
    for salle_data in salles:
        for serie_info in salle_data.get('series', []):
            sn = serie_info.get('nom')
            if sn and sn not in rec_pool:
                try:
                    so = Serie.objects.get(nom=sn, user=request.user)
                    inscs = Inscription.objects.filter(serie=so).order_by('num_ins')
                    rec_pool[sn] = [{'num_ins': ins.num_ins, 'nom_prenom': ins.nom_prenom, 'cin': ins.cin or ''} for ins in inscs]
                except Serie.DoesNotExist:
                    rec_pool[sn] = []

    for salle_data in salles:
        salle_numero = salle_data.get('salle_numero')
        section_str  = salle_data.get('section_str', '')
        series_str   = salle_data.get('series_str', '')
        candidats    = salle_data.get('candidats', [])

        if not candidats:
            for serie_info in salle_data.get('series', []):
                serie_nom = serie_info.get('nom')
                if not serie_nom: continue
                pool = rec_pool.get(serie_nom, [])
                if not pool:
                    errors.append(f"لا توجد أرقام متبقية للسلسلة {serie_nom}"); continue
                room_cap = int(salle_data.get('layout', 18))
                candidats = pool[:room_cap]
                del pool[:room_cap]
                if not series_str:
                    series_str = serie_nom
                if not section_str:
                    try:
                        so = Serie.objects.get(nom=serie_nom, user=request.user)
                        section_str = so.section.nom if so.section else section_str
                    except Serie.DoesNotExist:
                        pass
                break
            if not candidats:
                errors.append(f"Aucun candidat pour la salle {salle_numero}"); continue

        data = {
            'centre': centre,
            'annee_scolaire': annee,
            'delegation': request.user.delegation or '',
            'section': section_str,
            'serie': series_str,
            'salle': str(salle_numero),
            'matiere': matiere,
            'date': date,
        }

        try:
            stream = generate_rec_document(data, candidats)
            doc_id = str(uuid.uuid4())
            label = f"قائمة — قاعة {salle_numero} — {series_str} — {date}"
            generated_docs[doc_id] = {'stream': stream, 'label': label, 'type': 'rec'}
            documents.append({'doc_id': doc_id, 'label': label, 'type': 'rec'})
        except Exception as e:
            errors.append(f"قائمة élèves salle {salle_numero} : {str(e)}")

    if errors and not documents:
        return Response({"error": "\n".join(errors)}, status=400)

    response_data = {"documents": documents}
    if errors:
        response_data["warnings"] = errors
    return Response(response_data, status=200)


def merge_docx(streams):
    """Merge multiple DOCX streams into one, inserting content before final sectPr."""
    from copy import deepcopy
    from docx.oxml.ns import qn
    from lxml import etree

    streams[0].seek(0)
    merged = Document(streams[0])

    for stream in streams[1:]:
        stream.seek(0)
        doc = Document(stream)

        body = merged.element.body
        sectPr = body.find(qn('w:sectPr'))

        first = True
        for child in list(doc.element.body):
            if child.tag == qn('w:sectPr'):
                continue
            copied = deepcopy(child)
            if first and copied.tag == qn('w:p'):
                # Attach page break to the first paragraph instead of creating an orphan
                pPr = copied.find(qn('w:pPr'))
                if pPr is None:
                    pPr = etree.Element(qn('w:pPr'))
                    copied.insert(0, pPr)
                pageBrk = etree.Element(qn('w:pageBreakBefore'))
                pPr.insert(0, pageBrk)
                first = False
            if sectPr is not None:
                body.insert(body.index(sectPr), copied)
            else:
                body.append(copied)

    out = io.BytesIO()
    merged.save(out)
    out.seek(0)
    return out


TYPE_LABEL_MAP = {
    'plan': 'مخطط القاعة',
    'envelope': 'ملصق الظرف',
    'numero': 'أرقام المترشحين',
    'sortie': 'خروج طوارئ',
    'presence': 'بطاقات الحضور',
    'door': 'سجل القاعة',
}


def _merge_by_type(documents):
    """Merge all generated_docs by type — désactivé (merge_docx ne gère pas les images)."""
    return documents


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_verification(request):
    """
    Génère بطاقات التثبت (modele_rec.docx) — mergées en un seul fichier.
    """
    centre  = request.user.centre or ''
    annee   = request.user.annee_scolaire or ''
    deleg   = request.user.delegation or ''

    SECTION_ORDER = {
        'الآداب': 0,
        'علوم تجريبية': 1,
        'الاقتصاد و التصرف': 2,
        'علوم تقنية': 3,
        'علوم إعلامية': 4,
    }

    series_objs = sorted(
        Serie.objects.filter(user=request.user),
        key=lambda s: (SECTION_ORDER.get(s.section.nom if s.section else '', 999), s.nom)
    )
    if not series_objs:
        return Response({"error": "لا توجد سلاسل"}, status=400)

    streams = []
    labels = []
    errors = []
    salle_counter = 1

    for serie_obj in series_objs:
        inscs = Inscription.objects.filter(serie=serie_obj).order_by('num_ins')
        if not inscs.exists():
            continue

        candidats = [{'num_ins': ins.num_ins, 'nom_prenom': ins.nom_prenom, 'cin': ins.cin or ''} for ins in inscs]
        section_nom = serie_obj.section.nom if serie_obj.section else ''
        data = {
            'centre': centre,
            'annee_scolaire': annee,
            'delegation': deleg,
            'section': section_nom,
            'serie': serie_obj.nom,
            'salle': str(salle_counter),
            'matiere': '',
            'date': '',
        }
        salle_counter += 1
        try:
            stream = generate_rec_document(data, candidats)
            streams.append(stream)
            labels.append(f"{section_nom} — {serie_obj.nom}")
        except Exception as e:
            errors.append(f"{serie_obj.nom}: {str(e)}")

    if errors and not streams:
        return Response({"error": "\n".join(errors)}, status=400)

    try:
        if len(streams) == 1:
            merged = streams[0]
            label = f"بطاقات التثبت — {labels[0]}"
        else:
            # Return each series individually (merge désactivé: ne gère pas les images)
            docs_out = []
            for i, (s, lbl) in enumerate(zip(streams, labels)):
                did = str(uuid.uuid4())
                l = f"بطاقات التثبت — {lbl}"
                generated_docs[did] = {'stream': s, 'label': l, 'type': 'verif'}
                docs_out.append({'doc_id': did, 'label': l, 'type': 'verif'})
            response_data = {"documents": docs_out}
            if errors:
                response_data["warnings"] = errors
            return Response(response_data, status=200)
        doc_id = str(uuid.uuid4())
        generated_docs[doc_id] = {'stream': merged, 'label': label, 'type': 'verif'}

        response_data = {"documents": [{'doc_id': doc_id, 'label': label, 'type': 'verif'}]}
    except Exception as e:
        errors.append(f"دمج الوثائق: {str(e)}")
        if not errors:
            return Response({"error": "\n".join(errors)}, status=400)
        response_data = {"documents": [], "warnings": errors}
    else:
        if errors:
            response_data["warnings"] = errors

    return Response(response_data, status=200)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_sortie(request):
    """
    Génère le document sortie urgence (sortie_urgent.docx)
    pour les candidats de salles multiples avec sections multiples.
    """
    payload = request.data
    matiere     = payload.get('matiere', '')
    date        = payload.get('date', '')
    heure_debut = payload.get('heure_debut', '')
    heure_fin   = payload.get('heure_fin', '')
    centre      = payload.get('centre', '') or request.user.centre or ''
    delegation  = request.user.delegation or ''
    annee   = (request.user.annee_scolaire or '').split('/')[-1]  # "2025/2026" → "2026"
    salles_data = payload.get('salles', [])

    if not salles_data:
        return Response({"error": "لا توجد قاعات"}, status=400)

    documents = []
    errors = []

    for salle_data in salles_data:
        salle_numero = salle_data.get('salle_numero', '')
        sections_nom  = salle_data.get('sections', '')
        series_list  = salle_data.get('series', [])

        for serie_data in series_list:
            serie_nom     = serie_data.get('nom', '')
            candidats_data = serie_data.get('candidats', [])
            num_inscriptions = serie_data.get('num_inscriptions', [])

            if not candidats_data and not num_inscriptions:
                continue

            inscriptions_raw = []
            for cand in candidats_data:
                inscriptions_raw.append({
                    'num_ins': cand.get('num_ins', ''),
                    'nom_prenom': cand.get('nom_prenom', ''),
                    'section': cand.get('section', ''),
                })

            data = {
                'delegation':    delegation,
                'centre':        centre,
                'annee_scolaire': annee,
                'section':        sections_nom,
                'serie':          serie_nom,
                'salle':          str(salle_numero) if salle_numero else '',
                'matiere':        matiere,
                'date':           format_date_arabic(date),
                'heure_debut':    heure_debut,
                'heure_fin':      heure_fin,
            }

            try:
                stream = _generate_sortie(data, inscriptions_raw)
                doc_id = str(uuid.uuid4())
                label = f"خروج طوارئ — {sections_nom} — قاعة {salle_numero} — {serie_nom} — {date}" if salle_numero else f"خروج طوارئ — {sections_nom} — {serie_nom} — {date}"
                generated_docs[doc_id] = {'stream': stream, 'label': label, 'section': sections_nom, 'type': 'sortie'}
                documents.append({'doc_id': doc_id, 'label': label, 'section': sections_nom, 'type': 'sortie'})
            except Exception as e:
                errors.append(f"خروج {serie_nom}: {str(e)}")

    if errors and not documents:
        return Response({"error": "\n".join(errors)}, status=400)

    documents = _merge_by_type(documents)

    response_data = {"documents": documents}
    if errors:
        response_data["warnings"] = errors
    return Response(response_data, status=200)


# ==================== HELPERS ====================
