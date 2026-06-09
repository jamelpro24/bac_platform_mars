import random, io, os
from datetime import date, time
from collections import defaultdict

from django.db.models import Q
from django.http import HttpResponse

from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import Examen, ExamenSalle, Matiere, Salle, Session, SurveillancePlan, SurveillanceAssignment, SurveillanceGroupConfig, ProfessorSchedule
from professeurs.models import Professeur

MODELE_SUR_SEANCE = os.path.join(os.path.dirname(__file__), 'modeles', 'model_sur_seance.docx')
MODELE_RAPPORT_SURV = os.path.join(os.path.dirname(__file__), 'modeles', 'modele_rapport_surv.docx')

JOURS_ARABE = {0:'الإثنين', 1:'الثلاثاء', 2:'الأربعاء', 3:'الخميس', 4:'الجمعة', 5:'السبت', 6:'الأحد'}
MOIS_ARABE = {1:'يناير', 2:'فيفري', 3:'مارس', 4:'أفريل', 5:'ماي', 6:'جوان',
              7:'جويلية', 8:'أوت', 9:'سبتمبر', 10:'أكتوبر', 11:'نوفمبر', 12:'ديسمبر'}

def _format_date_arabic(date_obj):
    try:
        return f"{JOURS_ARABE[date_obj.weekday()]} {date_obj.day:02d} {MOIS_ARABE[date_obj.month]} {date_obj.year}"
    except Exception:
        return str(date_obj)




def _format_heure_arabe(t_str: str) -> str:
    """08:00 → 8 | 08:30 → 8 و30 | 09:15 → 9 و15"""
    if not t_str or ':' not in t_str:
        return t_str
    parts = t_str.split(':')
    h = parts[0].lstrip('0') or '0'
    m = parts[1]
    if m == '00':
        return h
    return f"{h} و{m}"

def _calc_heures(t_start: time, t_end: time) -> float:
    s = t_start.hour * 60 + t_start.minute
    e = t_end.hour * 60 + t_end.minute
    return round((e - s) / 60, 1)


def _get_salle_info(salle_id: int) -> tuple[int, str]:
    try:
        s = Salle.objects.get(id=salle_id)
        return s.numero, f"\u0642\u0627\u0639\u0629 {s.numero}"
    except Salle.DoesNotExist:
        return salle_id, f"\u0642\u0627\u0639\u0629 {salle_id}"


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_surveillance(request):
    session_id = request.data.get('session_id')
    if not session_id:
        return Response({"error": "معرف الدورة مطلوب"}, status=400)

    try:
        session = Session.objects.get(id=session_id, user=request.user)
    except Session.DoesNotExist:
        return Response({"error": "الدورة غير موجودة"}, status=404)

    examens_qs = Examen.objects.filter(session=session).order_by('date', 'heure_debut')
    if not examens_qs.exists():
        return Response({"error": "لا توجد امتحانات في هذه الدورة"}, status=400)

    all_dates = list(examens_qs.values_list('date', flat=True).distinct().order_by('date'))
    if len(all_dates) < 2:
        return Response({"error": "الدورة يجب أن تحتوي على يومين على الأقل"}, status=400)

    profs = list(Professeur.objects.filter(centre=request.user).order_by('nom'))
    if len(profs) < 4:
        return Response({"error": "يجب أن يكون هناك 4 أساتذة على الأقل"}, status=400)

    sur_hours: dict[int, float] = {p.id: 0.0 for p in profs}
    sup_hours: dict[int, float] = {p.id: 0.0 for p in profs}
    sur_counts: dict[int, int] = {p.id: 0 for p in profs}
    sup_counts: dict[int, int] = {p.id: 0 for p in profs}
    assignments_data: list[dict] = []

    def get_rooms_for_slot(day: date, t_start: time) -> set[int]:
        """Get all unique room IDs used for examens at this time on this day."""
        ex_ids = list(
            Examen.objects.filter(session=session, date=day, heure_debut=t_start)
            .values_list('id', flat=True)
        )
        room_ids = set()
        # From ExamenSalle (new model)
        for rid in ExamenSalle.objects.filter(examen_id__in=ex_ids).values_list('salle_id', flat=True):
            room_ids.add(rid)
        # From Examen.salle (old model)
        for e in Examen.objects.filter(id__in=ex_ids).exclude(salle__isnull=True).values_list('salle_id', flat=True):
            room_ids.add(e)
        # Reduce to 2/3 of rooms for surveillance
        room_list = sorted(room_ids)
        reduce_count = max(1, round(len(room_list) * 2 / 3))
        return set(room_list[:reduce_count])

    max_rooms_per_slot = 0
    seen_slots: set[tuple[date, time]] = set()
    for ex in examens_qs:
        slot = (ex.date, ex.heure_debut)
        if slot in seen_slots:
            continue
        seen_slots.add(slot)
        max_rooms_per_slot = max(max_rooms_per_slot, len(get_rooms_for_slot(ex.date, ex.heure_debut)))

    min_profs = max_rooms_per_slot * 3
    if min_profs and len(profs) < min_profs:
        return Response({
            "error": f"Nombre de professeurs insuffisant: il faut au moins {min_profs} professeurs pour couvrir {max_rooms_per_slot} salles avec 2 surveillants par salle et la liste de suppleance."
        }, status=400)

    def assign_block(block_dates: list[date], group_profs: list[Professeur], group_num: int):
        for day in block_dates:
            day_examens = examens_qs.filter(date=day).order_by('heure_debut')
            sessions: list[tuple[time, time, set[int]]] = []
            seen_times: set[tuple[int, int]] = set()
            for ex in day_examens:
                tkey = (ex.heure_debut.hour, ex.heure_debut.minute)
                if tkey in seen_times:
                    continue
                seen_times.add(tkey)
                # Find max end time for this time slot
                same_slot = day_examens.filter(heure_debut=ex.heure_debut)
                max_end = max(e.heure_fin for e in same_slot)
                rooms = get_rooms_for_slot(day, ex.heure_debut)
                sessions.append((ex.heure_debut, max_end, rooms))

            for session_idx, (t_start, t_end, room_ids) in enumerate(sessions, 1):
                if not room_ids:
                    continue
                needed = len(room_ids) * 2
                needed_suppleants = max(1, needed // 2)
                heures = _calc_heures(t_start, t_end)
                sorted_profs = sorted(group_profs, key=lambda p: (sur_counts[p.id], sur_hours[p.id], sup_counts[p.id], sup_hours[p.id], p.nom))
                assigned_ids: list[int] = []
                room_list = sorted(room_ids)
                room_slots = [(rid, slot) for rid in room_list for slot in range(2)]

                for prof, (rid, _slot) in zip(sorted_profs[:needed], room_slots):
                    sur_hours[prof.id] += heures
                    sur_counts[prof.id] += 1
                    salle_num, salle_label = _get_salle_info(rid)
                    assignments_data.append({
                        'professeur_id': prof.id,
                        'date': day,
                        'session_number': session_idx,
                        'time_start': t_start,
                        'time_end': t_end,
                        'salle_numero': salle_num,
                        'salle_label': salle_label,
                        'type': 'surveillant',
                        'group_number': group_num,
                        'heures': heures,
                    })
                    assigned_ids.append(prof.id)

                available_suppleants = [p for p in group_profs if p.id not in assigned_ids]
                sorted_suppleants = sorted(available_suppleants, key=lambda p: (sup_counts[p.id], sup_hours[p.id], sur_counts[p.id], sur_hours[p.id], p.nom))
                for prof in sorted_suppleants[:needed_suppleants]:
                    sup_hours[prof.id] += heures
                    sup_counts[prof.id] += 1
                    assignments_data.append({
                        'professeur_id': prof.id,
                        'date': day,
                        'session_number': session_idx,
                        'time_start': t_start,
                        'time_end': t_end,
                        'salle_numero': 0,
                        'salle_label': "\u0627\u062d\u062a\u064a\u0627\u0637",
                        'type': 'suppleant',
                        'group_number': group_num,
                        'heures': heures,
                    })

    assign_block(all_dates, profs, 1)

    plan = SurveillancePlan.objects.create(
        session=session,
        total_profs=len(profs),
        group_size=len(profs),
    )
    bulk = []
    for a in assignments_data:
        bulk.append(SurveillanceAssignment(
            plan=plan,
            professeur_id=a['professeur_id'],
            date=a['date'],
            session_number=a['session_number'],
            time_start=a['time_start'],
            time_end=a['time_end'],
            salle_numero=a['salle_numero'],
            salle_label=a['salle_label'],
            type=a['type'],
            group_number=a['group_number'],
            heures=a['heures'],
        ))
    SurveillanceAssignment.objects.bulk_create(bulk)

    prof_map = {p.id: p for p in profs}
    summary = [
        {
            'professeur_id': p.id,
            'nom': p.nom,
            'specialite': p.specialite,
            'sur_heures': sur_hours[p.id],
            'sup_heures': sup_hours[p.id],
            'sur_count': sur_counts[p.id],
            'sup_count': sup_counts[p.id],
            'total_heures': sur_hours[p.id] + sup_hours[p.id],
            'groupe': 1,
        }
        for p in profs
    ]

    return Response({
        'plan_id': plan.id,
        'total_profs': len(profs),
        'group_size': len(profs),
        'total_assignments': len(assignments_data),
        'summary': summary,
        'assignments': [
            {
                'id': idx,
                'professeur_nom': prof_map.get(a['professeur_id']).nom if prof_map.get(a['professeur_id']) else '?',
                'professeur_id': a['professeur_id'],
                'date': str(a['date']),
                'session_number': a['session_number'],
                'time_start': str(a['time_start']),
                'time_end': str(a['time_end']),
                'salle_numero': a['salle_numero'],
                'salle_label': a['salle_label'],
                'type': a['type'],
                'group_number': a['group_number'],
                'heures': a['heures'],
            }
            for idx, a in enumerate(assignments_data)
        ],
    })


def _normalize_spec(n: str) -> str:
    """Normalize a name for specialty matching (strip ال/و prefixes)."""
    import re
    words = re.sub(r'\s+', ' ', n.strip()).split()
    normalized = set()
    for w in words:
        while True:
            if w.startswith('ال') and len(w) > 2:
                w = w[2:]
            elif w.startswith('و'):
                w = w[1:]
            else:
                break
        if w:
            normalized.add(w)
    return ' '.join(sorted(normalized))


def _matiere_to_specs(nom: str) -> list[str]:
    """Convert a matiere name to possible specialty keywords."""
    n = _normalize_spec(nom)
    words = set(n.split())
    # Common aliases
    if 'عربية' in words or 'لغة' in words:
        return ['عربية', _normalize_spec('لغة عربية')]
    if 'فرنسية' in words or 'فرنسة' in words:
        return ['فرنسية']
    if 'انقليزية' in words or 'انجليز' in words:
        return ['انقليزية', 'انجليزية']
    if 'اعلامية' in words or 'إعلامية' in words:
        return ['اعلامية', 'إعلامية']
    if 'رياضيات' in words:
        return ['رياضيات']
    if all(w in words for w in ['علوم', 'حياة']) or all(w in words for w in ['علوم', 'ارض']):
        return ['علوم', _normalize_spec('علوم حياة')]
    if all(w in words for w in ['علوم', 'فيزي']) or 'فيزياء' in words:
        return ['فيزي', _normalize_spec('علوم فيزيائية')]
    if 'تاريخ' in words or 'جغرافيا' in words:
        return ['تاريخ', 'جغرافيا']
    if 'تقنية' in words or 'هندسة' in words:
        return ['تقنية', 'هندسة']
    if 'اقتصاد' in words or 'تصرف' in words:
        return ['اقتصاد', 'تصرف']
    if 'رياضة' in words or 'بدنية' in words or 'رياضي' in words:
        return ['رياضة', 'بدنية']
    if 'إسلامية' in words or 'تربية' in words:
        return ['إسلامية', 'تربية']
    if 'المانية' in words:
        return ['المانية']
    if 'اسبانية' in words or 'إسبانية' in words:
        return ['اسبانية', 'إسبانية']
    if 'تركية' in words:
        return ['تركية']
    if 'مدنية' in words:
        return ['مدنية']
    if 'تشكيلية' in words or 'فن' in words:
        return ['تشكيلية', 'فن']
    if 'موسيقية' in words or 'موسيق' in words:
        return ['موسيقية', 'موسيق']
    return [n]


def _can_supervise(prof_specialite: str, matiere_nom: str) -> bool:
    """Return True if the professor CAN supervise this subject (no conflict)."""
    pn = _normalize_spec(prof_specialite)
    mn = _normalize_spec(matiere_nom)
    if pn == mn:
        return False
    # Check word overlap
    p_words = set(pn.split())
    m_words = set(mn.split())
    if p_words & m_words:
        return False
    # Check aliases
    m_specs = _matiere_to_specs(matiere_nom)
    for ms in m_specs:
        msn = _normalize_spec(ms)
        if pn.find(msn) >= 0 or msn.find(pn) >= 0:
            return False
    return True


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_groups_surveillance(request):
    """Generate a fair surveillance schedule from two professor groups and save to DB."""
    session_id = request.data.get('session_id')
    group1_ids = request.data.get('group1_ids', [])
    group2_ids = request.data.get('group2_ids', [])
    is_controle = request.data.get('is_controle', False)
    controle_ids = request.data.get('controle_ids', [])
    if not session_id:
        return Response({"error": "معرف الدورة مطلوب"}, status=400)
    try:
        session = Session.objects.get(id=session_id, user=request.user)
    except Session.DoesNotExist:
        return Response({"error": "الدورة غير موجودة"}, status=404)

    all_prof_ids = list(group1_ids) + list(group2_ids) + list(controle_ids)
    profs = list(Professeur.objects.filter(id__in=all_prof_ids, centre=request.user))
    if len(profs) < 2:
        return Response({"error": "يجب أن يكون هناك أستاذان على الأقل"}, status=400)

    prof_map_by_id = {p.id: p for p in profs}

    examens_qs = Examen.objects.filter(session=session).order_by('date', 'heure_debut')
    if not examens_qs.exists():
        return Response({"error": "لا توجد امتحانات في هذه الدورة"}, status=400)

    all_dates = list(examens_qs.values_list('date', flat=True).distinct().order_by('date'))

    def get_all_rooms(day, t_start):
        ex_ids = list(
            Examen.objects.filter(session=session, date=day, heure_debut=t_start)
            .values_list('id', flat=True)
        )
        room_ids = set()
        for rid in ExamenSalle.objects.filter(examen_id__in=ex_ids).values_list('salle_id', flat=True):
            room_ids.add(rid)
        for e in Examen.objects.filter(id__in=ex_ids).exclude(salle__isnull=True).values_list('salle_id', flat=True):
            room_ids.add(e)
        # From candidat_assignments (optional subjects)
        for ex in Examen.objects.filter(id__in=ex_ids).exclude(candidat_assignments={}):
            ca = ex.candidat_assignments
            if isinstance(ca, dict) and 'rooms' in ca:
                for r in ca['rooms']:
                    if 'salle_id' in r:
                        room_ids.add(r['salle_id'])
        return sorted(room_ids)

    # Build slot structure first (all days, all slots)
    slot_structure: list[dict] = []
    for day in all_dates:
        day_examens = examens_qs.filter(date=day).order_by('heure_debut')
        seen_times: set[tuple[int, int]] = set()
        for ex in day_examens:
            tkey = (ex.heure_debut.hour, ex.heure_debut.minute)
            if tkey in seen_times:
                continue
            seen_times.add(tkey)
            same = day_examens.filter(heure_debut=ex.heure_debut)
            max_end = max(e.heure_fin for e in same)
            rooms = get_all_rooms(day, ex.heure_debut)
            if not rooms:
                continue
            # Pre-compute matiere for each room in this slot
            room_matieres: dict[int, str] = {}
            for rid in rooms:
                m = Examen.objects.filter(
                    session=session, date=day, heure_debut=ex.heure_debut
                ).filter(
                    affectations__salle_id=rid
                ).first()
                if not m:
                    m = Examen.objects.filter(
                        session=session, date=day, heure_debut=ex.heure_debut, salle_id=rid
                    ).first()
                if not m:
                    # Check candidat_assignments (optional subjects)
                    for ex_ca in Examen.objects.filter(
                        session=session, date=day, heure_debut=ex.heure_debut
                    ).exclude(candidat_assignments={}):
                        ca = ex_ca.candidat_assignments
                        if isinstance(ca, dict) and 'rooms' in ca:
                            for r in ca['rooms']:
                                if r.get('salle_id') == rid:
                                    m = ex_ca
                                    break
                        if m:
                            break
                if m:
                    room_matieres[rid] = m.matiere.nom
            slot_structure.append({
                'day': day,
                't_start': ex.heure_debut,
                't_end': max_end,
                'rooms': rooms,
                'room_matieres': room_matieres,
                'heures': _calc_heures(ex.heure_debut, max_end),
            })

    if not slot_structure:
        return Response({"error": "لا توجد قاعات مخصصة للامتحانات"}, status=400)

    sur_hours: dict[int, float] = {p.id: 0.0 for p in profs}
    sup_hours: dict[int, float] = {p.id: 0.0 for p in profs}
    sur_counts: dict[int, int] = {p.id: 0 for p in profs}
    sup_counts: dict[int, int] = {p.id: 0 for p in profs}
    assignments_data: list[dict] = []

    # Split dates: first half → group1, second half → group2
    split_idx = (len(all_dates) + 1) // 2
    group1_dates = set(all_dates[:split_idx])
    group2_dates = set(all_dates[split_idx:])

    # Group slots by day
    slot_structure.sort(key=lambda s: (s['day'], s['t_start']))
    slots_by_day: dict[date, list[dict]] = defaultdict(list)
    for slot in slot_structure:
        slots_by_day[slot['day']].append(slot)

    def _assign_pool_to_rooms(profs_pool: list[Professeur], rooms: list[int],
                              room_matieres: dict[int, str], day: date,
                              t_start: time, t_end: time, heures: float,
                              group_num: int, session_num: int,
                              enforce_gender_inst: bool = False) -> set[int]:
        """Assign profs per-room: each room independently picks 1st then 2nd surveillant,
        updating hours immediately after each pick for fairness.
        2nd pick prefers different institution + different sexe from 1st.
        If enforce_gender_inst=True, 2nd pick enforces opposite sex + different institution."""
        if not rooms:
            return set()
        used_ids: set[int] = set()
        room_list = sorted(rooms)

        def first_sort_key(p, rid):
            can = _can_supervise(p.specialite, room_matieres.get(rid, ''))
            total = sur_hours[p.id] + sup_hours[p.id]
            return (0 if can else 1, total, sur_hours[p.id], sup_hours[p.id], p.id)

        def second_sort_key(p, rid, first_prof):
            can = _can_supervise(p.specialite, room_matieres.get(rid, ''))
            total = sur_hours[p.id] + sup_hours[p.id]
            diff_inst = p.institution != first_prof.institution
            diff_sexe = bool(p.sexe and first_prof.sexe and p.sexe != first_prof.sexe)
            if can and diff_inst and diff_sexe:
                priority = 0
            elif can and diff_inst:
                priority = 1
            elif can:
                priority = 2
            elif diff_inst and diff_sexe:
                priority = 3
            elif diff_inst:
                priority = 4
            else:
                priority = 5
            return (priority, total, sur_hours[p.id], sup_hours[p.id], p.id)

        def record_assignment(prof, rid):
            salle_num, salle_label = _get_salle_info(rid)
            matiere_nom = room_matieres.get(rid, '')
            sur_hours[prof.id] += heures
            sur_counts[prof.id] += 1
            assignments_data.append({
                'professeur_id': prof.id,
                'date': day,
                'session_number': session_num,
                'time_start': t_start,
                'time_end': t_end,
                'salle_numero': salle_num,
                'salle_label': salle_label,
                'type': 'surveillant',
                'group_number': group_num,
                'heures': heures,
                'matiere_nom': matiere_nom,
            })

        for rid in room_list:
            remaining = [p for p in profs_pool if p.id not in used_ids]
            if not remaining:
                break

            # Pick 1st surveillant (can-supervise first, then lowest hours)
            remaining.sort(key=lambda p: first_sort_key(p, rid))
            prof1 = remaining[0]
            used_ids.add(prof1.id)
            record_assignment(prof1, rid)

            # Pick 2nd surveillant (prefer different institution + sexe)
            remaining2 = [p for p in profs_pool if p.id not in used_ids]
            if not remaining2:
                continue
            if enforce_gender_inst and prof1.sexe:
                candidates = [p for p in remaining2 if p.sexe and prof1.sexe
                              and p.sexe != prof1.sexe
                              and p.institution != prof1.institution]
                if candidates:
                    remaining2 = candidates
                else:
                    continue
            remaining2.sort(key=lambda p: second_sort_key(p, rid, prof1))
            prof2 = remaining2[0]
            used_ids.add(prof2.id)
            record_assignment(prof2, rid)

        return used_ids

    def _add_suppleants(pool: list[Professeur], needed_sup: int, day: date,
                        t_start: time, t_end: time, heures: float,
                        group_num: int, session_num: int) -> list[int]:
        """Pick needed_sup profs from pool (least hours first). Returns assigned IDs."""
        pool.sort(key=lambda p: (sur_hours[p.id] + sup_hours[p.id], sup_hours[p.id], sup_counts[p.id], p.id))
        assigned_ids = []
        for p in pool[:needed_sup]:
            sup_hours[p.id] += heures
            sup_counts[p.id] += 1
            assignments_data.append({
                'professeur_id': p.id,
                'date': day,
                'session_number': session_num,
                'time_start': t_start,
                'time_end': t_end,
                'salle_numero': 0,
                'salle_label': 'احتياط',
                'type': 'suppleant',
                'group_number': group_num,
                'heures': heures,
                'matiere_nom': '',
            })
            assigned_ids.append(p.id)
        return assigned_ids

    # --- Contrôle session: distribute g1+g2 profs evenly across days ---
    ctrl_day_prof_ids: dict[date, list[int]] = {}
    if is_controle:
        g1g2_ids = list(group1_ids) + list(group2_ids)
        random.shuffle(g1g2_ids)
        n_days = len(all_dates)
        per_day = len(g1g2_ids) // n_days
        rem = len(g1g2_ids) % n_days
        idx = 0
        for i, day in enumerate(all_dates):
            count = per_day + (1 if i < rem else 0)
            if count > 0:
                day_ids = list(g1g2_ids[idx:idx+count]) + list(controle_ids)
                ctrl_day_prof_ids[day] = day_ids
                idx += count

    for day in all_dates:
        day_slots = slots_by_day.get(day, [])
        if not day_slots:
            continue

        # Determine which group works this day
        if is_controle:
            day_profs = [p for p in profs if p.id in ctrl_day_prof_ids.get(day, [])]
            group_num = 0
            # Use ALL rooms for contrôle (not just exam-scheduled rooms)
            all_room_ids = list(Salle.objects.filter(user=request.user).values_list('id', flat=True).order_by('numero'))
            for slot in day_slots:
                slot['rooms'] = list(all_room_ids)
                slot['room_matieres'] = {rid: '' for rid in all_room_ids}
        elif day in group2_dates:
            day_profs = [p for p in profs if p.id in group2_ids]
            group_num = 2
        else:
            day_profs = [p for p in profs if p.id in group1_ids]
            group_num = 1

        if len(day_profs) < 2:
            continue

        day_surv_ids: set[int] = set()
        day_sup_ids: set[int] = set()

        for slot_idx, slot in enumerate(day_slots):
            t_start = slot['t_start']
            t_end = slot['t_end']
            heures = slot['heures']
            rooms = list(slot['rooms'])
            room_matieres = slot['room_matieres']

            if not rooms:
                continue

            if is_controle:
                # Each room needs 2 surv + 1 sup = 3 profs total
                max_rooms = len(day_profs) // 3
            else:
                max_rooms = len(day_profs) // 2
            if len(rooms) > max_rooms:
                rooms = rooms[:max_rooms]
            # Ensure at least 1 room
            if not rooms:
                rooms = [list(Salle.objects.filter(user=request.user).values_list('id', flat=True).order_by('numero'))[0]]

            needed_surv = len(rooms) * 2
            needed_sup = max(1, needed_surv // 2)

            if is_controle:
                # Phase 1: Suppleance first — guarantee every prof gets at least 1
                sup_pool = [p for p in day_profs if sup_counts[p.id] < 1]
                sup_pool.sort(key=lambda p: (sur_hours[p.id] + sup_hours[p.id], p.id))
                sup_taken = min(needed_sup, len(sup_pool))
                sup_ids_set = set(_add_suppleants(
                    sup_pool, sup_taken,
                    day, t_start, t_end, heures, 0, slot_idx + 1
                ))
                if sup_taken < needed_sup:
                    extra = [p for p in day_profs if p.id not in sup_ids_set]
                    extra.sort(key=lambda p: (sur_hours[p.id] + sup_hours[p.id], p.id))
                    sup_ids_set.update(_add_suppleants(
                        extra, needed_sup - sup_taken,
                        day, t_start, t_end, heures, 0, slot_idx + 1
                    ))
                # Phase 2: Surveillance from remaining profs
                surv_pool = [p for p in day_profs if p.id not in sup_ids_set]
                surv_pool.sort(key=lambda p: (sur_counts[p.id], sur_hours[p.id] + sup_hours[p.id], p.id))
                surv_taken = surv_pool[:needed_surv]
                surv_ids = _assign_pool_to_rooms(
                    surv_taken, rooms, room_matieres,
                    day, t_start, t_end, heures, 0, slot_idx + 1,
                    enforce_gender_inst=True
                )
                day_surv_ids.update(surv_ids)
                day_sup_ids.update(sup_ids_set)
            else:
                # Build pool: ALL sup>0 profs (any category) before any sup=0 profs
                # This ensures suppleance always picks from sup=0 first
                def pkey(p):
                    return (sur_hours[p.id] + sup_hours[p.id], p.id)
                c1a = [p for p in day_profs if p.id not in day_surv_ids and p.id not in day_sup_ids and sup_counts[p.id] > 0]
                c1b = [p for p in day_profs if p.id not in day_surv_ids and p.id not in day_sup_ids and sup_counts[p.id] == 0]
                c2a = [p for p in day_profs if p.id in day_sup_ids and sup_counts[p.id] > 0]
                c2b = [p for p in day_profs if p.id in day_sup_ids and sup_counts[p.id] == 0]
                c3a = [p for p in day_profs if p.id in day_surv_ids and sup_counts[p.id] > 0]
                c3b = [p for p in day_profs if p.id in day_surv_ids and sup_counts[p.id] == 0]
                c1a.sort(key=pkey); c1b.sort(key=pkey)
                c2a.sort(key=pkey); c2b.sort(key=pkey)
                c3a.sort(key=pkey); c3b.sort(key=pkey)
                pool = c1a + c2a + c3a + c1b + c2b + c3b
                pool = pool[:needed_surv + needed_sup]
                surv_ids = _assign_pool_to_rooms(
                    pool[:needed_surv], rooms, room_matieres,
                    day, t_start, t_end, heures, group_num, slot_idx + 1
                )
                sup_ids = _add_suppleants(
                    pool[needed_surv:], needed_sup,
                    day, t_start, t_end, heures, group_num, slot_idx + 1
                )
                day_surv_ids.update(surv_ids)
                day_sup_ids.update(sup_ids)

        # --- The daily guarantee and fallback blocks (previously here) are removed ---
        # to enforce the rule: suppleants = half of monitors, no extra.

    # --- Save to DB ---
    plan = SurveillancePlan.objects.create(
        session=session,
        total_profs=len(profs),
        group_size=len(profs),
    )
    bulk = []
    for a in assignments_data:
        bulk.append(SurveillanceAssignment(
            plan=plan,
            professeur_id=a['professeur_id'],
            date=a['date'],
            session_number=a['session_number'],
            time_start=a['time_start'],
            time_end=a['time_end'],
            salle_numero=a['salle_numero'],
            salle_label=a['salle_label'],
            type=a['type'],
            group_number=a.get('group_number', 1),
            heures=a['heures'],
        ))
    SurveillanceAssignment.objects.bulk_create(bulk)
    # Re-fetch with IDs (default order = insertion order = matches assignments_data)
    saved_assignments = list(SurveillanceAssignment.objects.filter(plan=plan))

    prof_list = profs
    summary = [
        {
            'professeur_id': p.id,
            'nom': p.nom,
            'specialite': p.specialite,
            'sur_heures': sur_hours[p.id],
            'sup_heures': sup_hours[p.id],
            'sur_count': sur_counts[p.id],
            'sup_count': sup_counts[p.id],
            'total_heures': sur_hours[p.id] + sup_hours[p.id],
            'groupe': 0 if is_controle else (1 if p.id in group1_ids else 2),
        }
        for p in prof_list
    ]

    return Response({
        'plan_id': plan.id,
        'total_assignments': len(assignments_data),
        'total_profs': len(profs),
        'summary': summary,
        'assignments': [
            {
                'id': sa.id,
                'professeur_nom': prof_map_by_id.get(a['professeur_id']).nom if prof_map_by_id.get(a['professeur_id']) else '?',
                'professeur_id': a['professeur_id'],
                'date': str(a['date']),
                'session_number': a['session_number'],
                'time_start': str(a['time_start']),
                'time_end': str(a['time_end']),
                'salle_numero': a['salle_numero'],
                'salle_label': a['salle_label'],
                'type': a['type'],
                'group_number': a['group_number'],
                'heures': a['heures'],
                'matiere_nom': a.get('matiere_nom', ''),
            }
            for a, sa in zip(assignments_data, saved_assignments)
        ],
    })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_surveillance_plans(request):
    plans = SurveillancePlan.objects.filter(session__user=request.user).order_by('-created_at')
    return Response([
        {
            'id': p.id,
            'session_id': p.session_id,
            'session_nom': p.session.nom,
            'total_profs': p.total_profs,
            'group_size': p.group_size,
            'created_at': p.created_at.isoformat(),
            'assignments_count': p.assignments.count(),
        }
        for p in plans
    ])


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_surveillance_plan_detail(request, plan_id):
    try:
        plan = SurveillancePlan.objects.get(id=plan_id, session__user=request.user)
    except SurveillancePlan.DoesNotExist:
        return Response({"error": "الخطة غير موجودة"}, status=404)

    assignments = plan.assignments.select_related('professeur').order_by('date', 'session_number', 'salle_numero')
    profs = Professeur.objects.filter(centre=request.user)
    hours_map: dict[int, float] = {}
    sur_hours_map: dict[int, float] = {}
    sup_hours_map: dict[int, float] = {}
    sur_counts_map: dict[int, int] = {}
    sup_counts_map: dict[int, int] = {}
    # Preload matiere for each (date, time_start, salle_numero)
    matiere_cache: dict[tuple, str] = {}
    def _get_matiere(day, t_start, salle_num):
        key = (day, t_start, salle_num)
        if key in matiere_cache:
            return matiere_cache[key]
        ex = Examen.objects.filter(
            session=plan.session, date=day, heure_debut=t_start
        ).filter(
            Q(affectations__salle_id__in=Salle.objects.filter(numero=salle_num).values('id'))
            | Q(salle_id__in=Salle.objects.filter(numero=salle_num).values('id'))
        ).first()
        if not ex:
            # Check candidat_assignments
            for ex_ca in Examen.objects.filter(session=plan.session, date=day, heure_debut=t_start).exclude(candidat_assignments={}):
                ca = ex_ca.candidat_assignments
                if isinstance(ca, dict) and 'rooms' in ca:
                    for r in ca['rooms']:
                        if r.get('salle_id') and Salle.objects.filter(id=r['salle_id'], numero=salle_num).exists():
                            ex = ex_ca
                            break
                if ex:
                    break
        result = ex.matiere.nom if ex else ''
        matiere_cache[key] = result
        return result

    for a in assignments:
        hours_map[a.professeur_id] = hours_map.get(a.professeur_id, 0) + a.heures
        if a.type == 'surveillant':
            sur_hours_map[a.professeur_id] = sur_hours_map.get(a.professeur_id, 0) + a.heures
            sur_counts_map[a.professeur_id] = sur_counts_map.get(a.professeur_id, 0) + 1
        else:
            sup_hours_map[a.professeur_id] = sup_hours_map.get(a.professeur_id, 0) + a.heures
            sup_counts_map[a.professeur_id] = sup_counts_map.get(a.professeur_id, 0) + 1

    # Load group config to determine group membership for saved plans
    try:
        config = SurveillanceGroupConfig.objects.get(session=plan.session)
        g1_ids = set(config.group1_ids)
        g2_ids = set(config.group2_ids)
    except SurveillanceGroupConfig.DoesNotExist:
        g1_ids = set()
        g2_ids = set()

    return Response({
        'plan_id': plan.id,
        'session_nom': plan.session.nom,
        'total_profs': plan.total_profs,
        'group_size': plan.group_size,
        'created_at': plan.created_at.isoformat(),
        'summary': [
            {
                'professeur_id': p.id,
                'nom': p.nom,
                'specialite': p.specialite,
                'institution': p.institution,
                'sexe': p.sexe,
                'sur_heures': sur_hours_map.get(p.id, 0),
                'sup_heures': sup_hours_map.get(p.id, 0),
                'sur_count': sur_counts_map.get(p.id, 0),
                'sup_count': sup_counts_map.get(p.id, 0),
                'total_heures': hours_map.get(p.id, 0),
                'groupe': 1 if p.id in g1_ids else (2 if p.id in g2_ids else None),
            }
            for p in profs
        ],
        'assignments': [
            {
                'id': a.id,
                'professeur_nom': a.professeur.nom,
                'professeur_id': a.professeur_id,
                'professeur_institution': a.professeur.institution,
                'professeur_sexe': a.professeur.sexe,
                'date': str(a.date),
                'session_number': a.session_number,
                'time_start': str(a.time_start),
                'time_end': str(a.time_end),
                'salle_numero': a.salle_numero,
                'salle_label': a.salle_label,
                'type': a.type,
                'group_number': a.group_number,
                'heures': a.heures,
                'matiere_nom': _get_matiere(a.date, a.time_start, a.salle_numero) if a.time_start else '',
            }
            for a in assignments
        ],
    })


@api_view(['PUT'])
@permission_classes([IsAuthenticated])
def update_surveillance_assignment(request, assignment_id):
    """Update a single SurveillanceAssignment (e.g. swap professor)."""
    try:
        assignment = SurveillanceAssignment.objects.get(id=assignment_id, plan__session__user=request.user)
    except SurveillanceAssignment.DoesNotExist:
        return Response({"error": "التعيين غير موجود"}, status=404)

    professeur_id = request.data.get('professeur_id')
    if professeur_id is not None:
        try:
            prof = Professeur.objects.get(id=professeur_id, centre=request.user)
        except Professeur.DoesNotExist:
            return Response({"error": "الأستاذ غير موجود"}, status=404)
        assignment.professeur = prof

    if 'salle_numero' in request.data:
        assignment.salle_numero = request.data['salle_numero']
    if 'salle_label' in request.data:
        assignment.salle_label = request.data['salle_label']
    if 'type' in request.data:
        assignment.type = request.data['type']

    assignment.save()
    return Response({
        'id': assignment.id,
        'professeur_id': assignment.professeur_id,
        'professeur_nom': assignment.professeur.nom,
        'date': str(assignment.date),
        'session_number': assignment.session_number,
        'time_start': str(assignment.time_start),
        'time_end': str(assignment.time_end),
        'salle_numero': assignment.salle_numero,
        'salle_label': assignment.salle_label,
        'type': assignment.type,
        'group_number': assignment.group_number,
        'heures': assignment.heures,
    })


@api_view(['GET', 'PUT'])
@permission_classes([IsAuthenticated])
def surveillance_groups(request):
    session_id = request.query_params.get('session_id') or request.data.get('session_id')
    if not session_id:
        return Response({"error": "session_id مطلوب"}, status=400)
    try:
        session = Session.objects.get(id=session_id, user=request.user)
    except Session.DoesNotExist:
        return Response({"error": "الدورة غير موجودة"}, status=404)

    if request.method == 'GET':
        try:
            config = SurveillanceGroupConfig.objects.get(session=session)
            return Response({
                'group1_ids': config.group1_ids,
                'group2_ids': config.group2_ids,
            })
        except SurveillanceGroupConfig.DoesNotExist:
            return Response({'group1_ids': [], 'group2_ids': []})

    # PUT
    group1_ids = request.data.get('group1_ids', [])
    group2_ids = request.data.get('group2_ids', [])
    config, _ = SurveillanceGroupConfig.objects.update_or_create(
        session=session,
        defaults={'group1_ids': group1_ids, 'group2_ids': group2_ids},
    )
    return Response({
        'group1_ids': config.group1_ids,
        'group2_ids': config.group2_ids,
    })


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def download_surveillance_doc(request):
    """Generate a Word document for a day+session (all rooms) using model_sur_seance.docx template."""
    plan_id = request.data.get('plan_id')
    doc_date = request.data.get('date')
    session_number = request.data.get('session_number')

    if not all([plan_id, doc_date, session_number]):
        return Response({"error": "plan_id, date, session_number مطلوبة"}, status=400)

    try:
        plan = SurveillancePlan.objects.get(id=plan_id, session__user=request.user)
    except SurveillancePlan.DoesNotExist:
        return Response({"error": "الخطة غير موجودة"}, status=404)

    # Get all surveillance assignments for this day+session
    all_assignments = SurveillanceAssignment.objects.filter(
        plan=plan, date=doc_date, session_number=session_number
    ).select_related('professeur').order_by('salle_numero', 'id')

    # Group by room
    from collections import defaultdict
    rooms: dict[int, list] = defaultdict(list)
    suppleants: list = []
    for a in all_assignments:
        if a.type == 'surveillant' and a.salle_numero > 0:
            rooms[a.salle_numero].append(a)
        elif a.type == 'suppleant':
            suppleants.append(a)

    # Get session time from the first assignment (authoritative source)
    heure_debut = ''
    heure_fin = ''
    if all_assignments.exists():
        first = all_assignments.first()
        if first.time_start:
            heure_debut = str(first.time_start)[:5]
        if first.time_end:
            heure_fin = str(first.time_end)[:5]

    centre = request.user.centre or ''
    delegation = request.user.delegation or ''
    annee_scolaire = request.user.annee_scolaire or ''

    try:
        d = date.fromisoformat(doc_date)
        date_ar = _format_date_arabic(d)
    except Exception:
        date_ar = doc_date

    try:
        doc = Document(MODELE_SUR_SEANCE)
    except Exception:
        return Response({"error": "قالب الوثيقة غير موجود"}, status=500)

    # Global placeholders (session-wide) — only for header tables
    _replace_placeholder_docx(doc, 'annee_scolaire', annee_scolaire)
    _replace_placeholder_docx(doc, 'centre', centre)
    _replace_placeholder_docx(doc, 'delegation', delegation)
    _replace_placeholder_docx(doc, 'municipalité', delegation)
    _replace_placeholder_docx(doc, 'date', date_ar)
    _replace_placeholder_docx(doc, 'seance', str(session_number))
    _replace_placeholder_docx(doc, 'temp', _format_heure_arabe(heure_debut))
    _replace_placeholder_docx(doc, 'specialité', '')

    # Pre-fetch exam data per room (serie, section, matiere, heure)
    session_time_start = all_assignments.first().time_start if all_assignments.exists() else None
    room_exam_data: dict[int, dict] = {}

    def _build_room_exam_data(time_filter):
        """Build room_exam_data dict, optionally filtered by heure_debut."""
        q = Q(examen__date=doc_date) & Q(examen__user=request.user)
        if time_filter and session_time_start:
            q &= Q(examen__heure_debut=session_time_start)
        out: dict[int, dict] = {}
        for es in ExamenSalle.objects.filter(q).select_related(
            'examen', 'examen__matiere', 'examen__serie', 'examen__section', 'salle'
        ):
            s_num = es.salle.numero
            ex = es.examen
            if s_num not in out:
                out[s_num] = {
                    'sections': [],
                    'matieres': [],
                    'series': [],
                    'heure_debut': str(ex.heure_debut)[:5] if ex.heure_debut else '',
                    'heure_fin': str(ex.heure_fin)[:5] if ex.heure_fin else '',
                }
            sec = ex.section.nom if ex.section else ''
            mat = ex.matiere.nom if ex.matiere else ''
            ser = ex.serie.nom if ex.serie else ''
            if sec and sec not in out[s_num]['sections']:
                out[s_num]['sections'].append(sec)
            if mat and mat not in out[s_num]['matieres']:
                out[s_num]['matieres'].append(mat)
            if ser and ser not in out[s_num]['series']:
                out[s_num]['series'].append(ser)
        # Collapse lists to display strings
        for v in out.values():
            v['section'] = '\n'.join(v['sections'])
            v['matiere'] = '\n'.join(v['matieres'])
            v['serie'] = '\n'.join(v['series'])
        return out

    # First try with session time filter; if a room has no data, fall back to unfiltered
    # and override its time with the session's actual time.
    room_exam_data = _build_room_exam_data(True)
    if room_exam_data:
        fallback_all = _build_room_exam_data(False)
        for s_num in fallback_all:
            if s_num not in room_exam_data:
                d = fallback_all[s_num]
                d['heure_debut'] = heure_debut
                d['heure_fin'] = heure_fin
                room_exam_data[s_num] = d
    else:
        room_exam_data = _build_room_exam_data(False)
        for d in room_exam_data.values():
            d['heure_debut'] = heure_debut
            d['heure_fin'] = heure_fin

    # Fill Table 1: one row per room
    if len(doc.tables) >= 2:
        table1 = doc.tables[1]
        room_nums = sorted(rooms.keys())
        from copy import deepcopy
        for i, rn in enumerate(room_nums):
            row_idx = i + 1
            if row_idx >= len(table1.rows):
                new_row = deepcopy(table1.rows[1]._element)
                table1._element.append(new_row)
            row = table1.rows[row_idx]
            profs = rooms[rn]
            ed = room_exam_data.get(rn, {})
            salle_label = profs[0].salle_label if hasattr(profs[0], 'salle_label') and profs[0].salle_label else f"قاعة {rn}"

            def _set_cell(p, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10), color=None):
                p.clear()
                p.alignment = align
                r = p.add_run(text)
                r.font.size = size
                r.font.bold = bold
                if color:
                    r.font.color.rgb = color
                return r

            _set_cell(row.cells[0].paragraphs[0], ed.get('section', ''))
            _set_cell(row.cells[1].paragraphs[0], ed.get('matiere', ''))
            _set_cell(row.cells[2].paragraphs[0],
                      f"{_format_heure_arabe(ed.get('heure_debut', ''))}-{_format_heure_arabe(ed.get('heure_fin', ''))}",
                      align=WD_ALIGN_PARAGRAPH.RIGHT)
            # Set RTL on time cell
            p2 = row.cells[2].paragraphs[0]
            pPr = p2._element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)

            _set_cell(row.cells[3].paragraphs[0], salle_label)

            # Cell 4: surveillant 1 — nom + spécialité + établissement
            row.cells[4].paragraphs[0].clear()
            p4 = row.cells[4].paragraphs[0]
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rn = p4.add_run(profs[0].professeur.nom)
            rn.font.size = Pt(10)
            rn.font.bold = True
            rs = p4.add_run(f"\n{profs[0].professeur.specialite or ''}")
            rs.font.size = Pt(8)
            rs.font.color.rgb = RGBColor(100, 100, 100)
            ri = p4.add_run(f"\n{profs[0].professeur.institution or ''}")
            ri.font.size = Pt(8)
            ri.font.color.rgb = RGBColor(100, 100, 100)

            # Cell 5: surveillant 2 — nom + spécialité + établissement
            row.cells[5].paragraphs[0].clear()
            p5 = row.cells[5].paragraphs[0]
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(profs) > 1:
                rn2 = p5.add_run(profs[1].professeur.nom)
                rn2.font.size = Pt(10)
                rn2.font.bold = True
                rs2 = p5.add_run(f"\n{profs[1].professeur.specialite or ''}")
                rs2.font.size = Pt(8)
                rs2.font.color.rgb = RGBColor(100, 100, 100)
                ri2 = p5.add_run(f"\n{profs[1].professeur.institution or ''}")
                ri2.font.size = Pt(8)
                ri2.font.color.rgb = RGBColor(100, 100, 100)
            else:
                r3 = p5.add_run('---')
                r3.font.bold = True

    # Fill Table 2: suppleance
    if len(doc.tables) >= 3:
        table2 = doc.tables[2]
        for i, a in enumerate(suppleants):
            row_idx = i + 4
            if row_idx >= len(table2.rows):
                from copy import deepcopy as _dc
                new_row = _dc(table2.rows[-1]._element)
                table2._element.append(new_row)
            row = table2.rows[row_idx]
            row.cells[0].paragraphs[0].clear()
            row.cells[0].paragraphs[0].add_run(str(i + 1))
            row.cells[1].paragraphs[0].clear()
            row.cells[1].paragraphs[0].add_run(_format_heure_arabe(heure_debut))
            row.cells[2].paragraphs[0].clear()
            row.cells[2].paragraphs[0].add_run(a.professeur.nom)
            row.cells[3].paragraphs[0].clear()
            row.cells[3].paragraphs[0].add_run(a.professeur.specialite or '')

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    response = HttpResponse(
        stream.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    filename = f"حراسة_{doc_date}_حصة{session_number}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def confirm_surveillance_schedule(request):
    """Flatten current SurveillanceAssignments into ProfessorSchedule (1 row per prof/day)."""
    plan_id = request.data.get('plan_id')
    ctrl_plan_id = request.data.get('ctrl_plan_id')

    plan_ids = []
    if plan_id:
        plan_ids.append(int(plan_id))
    if ctrl_plan_id:
        plan_ids.append(int(ctrl_plan_id))

    if not plan_ids:
        return Response({"error": "plan_id مطلوب"}, status=400)

    qs = SurveillanceAssignment.objects.filter(
        plan_id__in=plan_ids,
        plan__session__user=request.user,
    ).select_related('professeur')

    if not qs.exists():
        return Response({"error": "لا توجد تعيينات لهذه الخطة"}, status=404)

    # Group by (professeur_id, date)
    from collections import defaultdict
    groups = defaultdict(lambda: {'s1': None, 's2': None})
    for a in qs:
        key = (a.professeur_id, a.date)
        if a.session_number == 1:
            groups[key]['s1'] = a
        elif a.session_number == 2:
            groups[key]['s2'] = a

    # Delete all existing ProfessorSchedule for these profs
    prof_ids = list(set(k[0] for k in groups.keys()))
    ProfessorSchedule.objects.filter(professeur_id__in=prof_ids).delete()

    # Insert new rows
    new_rows = []
    for (pid, d), slots in groups.items():
        s1 = slots['s1']
        s2 = slots['s2']
        new_rows.append(ProfessorSchedule(
            professeur_id=pid,
            date=d,
            session1_type=s1.type if s1 else '',
            session1_start=s1.time_start if s1 else None,
            session1_end=s1.time_end if s1 else None,
            session1_salle=s1.salle_numero if s1 else None,
            session2_type=s2.type if s2 else '',
            session2_start=s2.time_start if s2 else None,
            session2_end=s2.time_end if s2 else None,
            session2_salle=s2.salle_numero if s2 else None,
        ))
    ProfessorSchedule.objects.bulk_create(new_rows)

    return Response({"success": True, "count": len(new_rows)})


# ==================== RAPPORT GROUPES SURVEILLANCE ====================

def _set_cell_surv(cell, text, size=10, bold=False):
    from docx.oxml.ns import qn
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Tajawal'
    # RTL
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Tajawal')


def _replace_placeholder_docx(doc, placeholder, value):
    """Replace {{placeholder}} in all paragraphs and table cells (handles split runs)."""
    from docx.oxml.ns import qn
    tag = '{{' + placeholder + '}}'
    val = str(value)

    def replace_in_para(para):
        runs = list(para._element.iter(qn('w:r')))
        # Build full text with position tracking
        parts = []  # (run, t_elem, start_pos, text)
        pos = 0
        for r in runs:
            for t in r.iter(qn('w:t')):
                txt = t.text or ''
                parts.append((r, t, pos, txt))
                pos += len(txt)
        full = ''.join(p[3] for p in parts)
        if tag not in full:
            return

        idx = full.index(tag)
        end = idx + len(tag)
        placed = False
        for r, t, pos, txt in parts:
            t_start = pos
            t_end = pos + len(txt)
            if t_end <= idx or t_start >= end:
                continue  # completely outside tag
            if t_start < idx:
                # starts before tag → keep left portion
                t.text = txt[:idx - t_start]
            elif t_end <= end:
                # entirely inside tag → clear
                t.text = ''
            else:
                # starts inside tag, extends beyond → keep right portion
                t.text = txt[end - t_start:]
            if not placed and t_start <= idx < t_end:
                # first overlapping run → insert value
                t.text = (t.text or '') + val
                placed = True

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_surv_report(request):
    """Generate professor hour report DOCX (totals from all sessions)."""
    plan_id = request.data.get('plan_id')
    ctrl_plan_id = request.data.get('ctrl_plan_id')

    plan_ids = [p for p in [plan_id, ctrl_plan_id] if p]
    if not plan_ids:
        return Response({"error": "plan_id requis"}, status=400)

    assignments = SurveillanceAssignment.objects.filter(plan_id__in=plan_ids)
    if not assignments.exists():
        return Response({"error": "aucune assignation trouvée"}, status=400)

    # Aggregate per professor across ALL groups (principale + contrôle)
    from collections import defaultdict
    prof_data = defaultdict(lambda: {'surv': 0.0, 'supp': 0.0, 'nom': ''})

    for a in assignments.select_related('professeur'):
        pid = a.professeur_id
        entry = prof_data[pid]
        entry['nom'] = a.professeur.nom or ''
        if a.type == 'suppleant':
            entry['supp'] += float(a.heures or 0)
        else:
            entry['surv'] += float(a.heures or 0)

    entries = sorted(prof_data.values(), key=lambda x: x['nom'])
    for e in entries:
        e['total'] = round(e['surv'] + e['supp'], 1)

    if not entries:
        return Response({"error": "aucune donnée"}, status=400)

    centre = request.user.centre or ''
    delegation = request.user.delegation or ''
    annee = request.user.annee_scolaire or ''

    from docx import Document as DocxDoc
    import io

    doc = DocxDoc(MODELE_RAPPORT_SURV)
    _replace_placeholder_docx(doc, 'delegation', delegation)
    _replace_placeholder_docx(doc, 'centre', centre)
    _replace_placeholder_docx(doc, 'annee_scolaire', annee)
    tab = doc.tables[0]
    while len(tab.rows) > 2:
        _delete_row_surv(tab, len(tab.rows) - 1)

    for i, e in enumerate(entries):
        if i == 0:
            row = tab.rows[1]
            for cell in row.cells:
                cell.paragraphs[0].clear()
        else:
            row = _clone_row_surv(tab, 1)

        # Template LTR order: total, supp, surv, name
        _set_cell_surv(row.cells[0], str(e['total']), size=10, bold=True)
        _set_cell_surv(row.cells[1], str(e['supp']), size=10)
        _set_cell_surv(row.cells[2], str(e['surv']), size=10)
        _set_cell_surv(row.cells[3], e['nom'], size=10)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    from .generate_views import generated_docs
    import uuid
    doc_id = str(uuid.uuid4())
    generated_docs[doc_id] = {'stream': stream, 'label': 'rapport_heures', 'type': 'surv_report'}
    result = [{'doc_id': doc_id, 'label': 'تقرير الحصص', 'type': 'surv_report'}]

    return Response({"documents": result})


def _delete_row_surv(table, row_idx):
    table._tbl.remove(table.rows[row_idx]._tr)


def _clone_row_surv(table, source_row_idx):
    import copy
    new_tr = copy.deepcopy(table.rows[source_row_idx]._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]
